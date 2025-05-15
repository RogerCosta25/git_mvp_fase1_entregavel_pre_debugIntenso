"""
Processador de documentos DOCX para o sistema de peticionamento.

Este módulo contém a classe DocumentoProcessor, responsável por processar
documentos DOCX, substituindo campos e aplicando regras condicionais.
"""

import os
import re
import json
import docx
import time
from typing import Dict, List, Any, Optional, Tuple, Set
from docx.document import Document  # Para tipagem
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

import locale
from decimal import Decimal

from src.motor_regras import MotorRegras
from src.exceptions import (
    ProcessamentoDocumentoError,
    TemplateNaoEncontradoError,
    TemplateError,
    DadosError
)
from src.logger import logger

class DocumentoProcessor:
    """
    Processador de documentos DOCX.
    """
    
    def __init__(self, motor_regras: Optional[MotorRegras] = None, modo_estrito: bool = False):
        """
        Inicializa o processador de documentos.
        
        Args:
            motor_regras: Instância do MotorRegras para processamento condicional.
                         Se None, cria uma nova instância.
            modo_estrito: Se True, valida todos os campos mesmo que não usados.
        """
        self.motor_regras = motor_regras or MotorRegras(usar_modelo_relacional=True)
        self.modo_estrito = modo_estrito
        self.campos_encontrados = set()
        self.campos_substituidos = set()
        self.campos_ausentes = set()
        self.campos_obrigatorios_ausentes = set()
        self.estatisticas_processamento = {}
        self.secoes_ativas = []
        self.secoes_encontradas = set()
        self.secoes_identificadas_pelo_conteudo = {}
        
        # Mapeamento padronizado entre seções e campos condicionais
        self.mapeamento_secoes = {
            'HORAS_EXTRAS': {
                'campo_condicional': 'calcular_horas_extras',
                'valor_ativo': 'Sim',
                'palavras_chave': ['hora extra', 'horas extras', 'jornada', 'sobrejornada'],
                'descricao': 'Seção de horas extras'
            },
            'VERBAS_RESCISORIAS': {
                'campo_condicional': 'motivo_rescisao',
                'valor_ativo': 'demissão sem justa causa',
                'palavras_chave': ['verba rescisória', 'verbas rescisórias', 'rescisão', 'verbas'],
                'descricao': 'Seção de verbas rescisórias'
            },
            'aviso_previo': {
                'campo_condicional': 'dias_aviso_previo_base_calculo',
                'valor_ativo': '30 dias',
                'palavras_chave': ['aviso prévio', 'aviso-prévio'],
                'descricao': 'Seção de aviso prévio'
            },
            'acumulo_funcao': {
                'campo_condicional': 'calcular_acumulo_funcao',
                'valor_ativo': 'Sim',
                'palavras_chave': ['acúmulo de função', 'acumulo de funcao', 'desvio de função'],
                'descricao': 'Seção de acúmulo de função'
            },
            'INSALUBRIDADE': {
                'campo_condicional': 'calcular_insalubridade',
                'valor_ativo': 'Sim',
                'palavras_chave': ['insalubridade', 'ambiente insalubre', 'adicional de insalubridade'],
                'descricao': 'Seção de insalubridade'
            }
        }
    
    def processar_documento(self, 
                          template_path: str, 
                          dados: Dict[str, Any], 
                          secoes_ativas: List[str], 
                          output_path: str) -> str:
        """
        Processa um documento DOCX, substituindo campos e aplicando regras condicionais.
        
        Args:
            template_path: Caminho para o arquivo de template DOCX.
            dados: Dicionário com os dados a serem inseridos no documento.
            secoes_ativas: Lista de IDs das seções que devem estar ativas.
            output_path: Caminho para salvar o documento processado.
            
        Returns:
            Caminho do documento gerado.
            
        Raises:
            TemplateError: Se o template não puder ser aberto ou processado.
            DadosError: Se os dados estiverem incompletos ou inválidos.
        """
        # Reinicia contadores e estatísticas para este processamento
        self.campos_encontrados = set()
        self.campos_substituidos = set()
        self.campos_ausentes = set()
        self.campos_obrigatorios_ausentes = set()
        self.estatisticas_processamento = {}
        self.secoes_encontradas = set()
        
        # Registra a lista de seções ativas (pode ser vazia e determinada automaticamente pelo motor)
        self.secoes_ativas = secoes_ativas
        
        tempo_inicio = time.time()
        
        # 1. Carregar o template
        try:
            # Verificar se o arquivo existe
            if not os.path.exists(template_path):
                mensagem = f"Template não encontrado: {template_path}"
                logger.error(mensagem)
                raise TemplateNaoEncontradoError(mensagem)
                
            # Corrigido: Usar docx.Document em vez de Document diretamente
            doc = docx.Document(template_path)
            logger.info(f"Template carregado: {template_path}")
        except Exception as e:
            mensagem = f"Erro ao abrir o template: {str(e)}"
            logger.error(mensagem)
            raise TemplateError(mensagem) from e
        
        # 2. Se não há seções ativas especificadas, determina automaticamente
        if not self.secoes_ativas:
            self.secoes_ativas = self._determinar_secoes_ativas(dados)
            logger.info(f"Seções ativas determinadas automaticamente: {self.secoes_ativas}")
        
        # 3. Substitui campos no documento
        doc = self._substituir_todos_campos(doc, dados)
            
        # 4. Processa seções condicionais
        doc = self._processar_secoes_condicionais(doc, dados)
            
        # 5. Realiza verificações finais e validações
        self._validar_documento(dados)
            
        # 6. Salva o documento processado
        try:
            # Garante que o diretório existe
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logger.info(f"Documento processado salvo em: {output_path}")
        except Exception as e:
            mensagem = f"Erro ao salvar o documento processado: {str(e)}"
            logger.error(mensagem)
            raise TemplateError(mensagem) from e
            
        # 7. Coleta estatísticas
        tempo_total = time.time() - tempo_inicio
        self.estatisticas_processamento = {
            'campos_encontrados': len(self.campos_encontrados),
            'campos_substituidos': len(self.campos_substituidos),
            'campos_ausentes': len(self.campos_ausentes),
            'campos_obrigatorios_ausentes': len(self.campos_obrigatorios_ausentes),
            'secoes_encontradas': len(self.secoes_encontradas),
            'secoes_ativas': len(self.secoes_ativas),
            'tempo_processamento': f"{tempo_total:.3f} segundos"
        }
        
        logger.info(f"Estatísticas do processamento: {self.estatisticas_processamento}")
        
        return output_path
    
    def _processar_tabela(self, tabela: Table, dados: Dict[str, Any]) -> None:
        """
        Processa todos os parágrafos em uma tabela.
        
        Args:
            tabela: Tabela a ser processada.
            dados: Dicionário com os dados para substituição.
        """
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    self._processar_paragrafo(paragrafo, dados)
    
    def _processar_secoes_condicionais(self, doc: Document, dados: Dict[str, Any]) -> Document:
        """
        Identifica e processa seções condicionais no documento.
        
        Args:
            doc: Documento a ser processado.
            dados: Dicionário com os dados para substituição.
            
        Returns:
            Documento com seções condicionais processadas.
        """
        # Se não há seções ativas definidas, faz a determinação automática
        if not self.secoes_ativas:
            self.secoes_ativas = self._determinar_secoes_ativas(dados)
            
        # Se ainda não há seções ativas, retorna o documento original
        if not self.secoes_ativas:
            logger.warning("Nenhuma seção ativa definida. O documento será processado sem remoção de seções.")
            return doc
            
        logger.info(f"Processando seções condicionais. Seções ativas: {self.secoes_ativas}")
        
        # Mapeia as seções no documento (por marcadores ou títulos)
        secoes_no_documento = self._mapear_secoes_no_documento(doc)
        
        # Se não encontrou seções, não há o que processar
        if not secoes_no_documento:
            logger.warning("Nenhuma seção encontrada no documento. O documento será processado sem remoção de seções.")
            return doc
        
        # Loga as seções encontradas para diagnóstico
        logger.info(f"Seções encontradas no documento: {list(secoes_no_documento.keys())}")
        self.secoes_encontradas = set(secoes_no_documento.keys())
        
        # Processa cada seção do documento
        for secao_id, info in secoes_no_documento.items():
            # Determina se a seção deve estar ativa
            ativa = secao_id in self.secoes_ativas
            
            logger.info(f"Seção '{secao_id}': {'ATIVA' if ativa else 'INATIVA'}")
            
            # Remove as seções inativas
            if not ativa:
                # Opção 1: Remover todos os elementos da seção
                if 'elementos' in info and info['elementos']:
                    for elemento in info['elementos']:
                        # Método seguro para "remover" elementos
                        if isinstance(elemento, Paragraph):
                            # Limpa o texto do parágrafo
                            logger.debug(f"Limpando conteúdo do parágrafo: '{elemento.text}'")
                            elemento.text = ""
                            
                            # Também limpa os runs
                            for run in elemento.runs:
                                run.text = ""
                
                logger.info(f"Seção '{secao_id}' removida (inativa)")
        
        # Realiza limpeza final do documento (remove parágrafos vazios)
        doc = self._limpar_documento_apos_processamento(doc)
        
        return doc
    
    def _limpar_documento_apos_processamento(self, doc: Document) -> Document:
        """
        Limpa o documento após processamento, removendo elementos vazios e trechos de marcação.
        
        Args:
            doc: Documento a ser limpo.
            
        Returns:
            Documento limpo.
        """
        logger.info("Realizando limpeza final do documento")
        
        # Processa todos os parágrafos para remover marcadores e parágrafos vazios
        paragrafos_para_remover = []
        
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text.strip()
            
            # Verifica se contém apenas marcadores de seção
            if re.match(r'^{{[\s]*[#/].*}}$', texto):
                logger.debug(f"Marcador de seção encontrado para remoção: '{texto}'")
                paragrafos_para_remover.append(paragrafo)
                continue
                
            # Verifica se está vazio (após processamento de seção inativa)
            if not texto and paragrafo.runs:
                # Verifica se todos os runs estão vazios
                todos_vazios = True
                for run in paragrafo.runs:
                    if run.text.strip():
                        todos_vazios = False
                        break
                
                if todos_vazios:
                    logger.debug(f"Parágrafo vazio encontrado para remoção (índice {i})")
                    paragrafos_para_remover.append(paragrafo)
        
        # Remove os parágrafos identificados
        # Nota: não podemos remover diretamente durante a iteração
        for paragrafo in paragrafos_para_remover:
            try:
                p = paragrafo._element
                p.getparent().remove(p)
            except Exception as e:
                logger.error(f"Erro ao remover parágrafo: {str(e)}")
        
        logger.info(f"Removidos {len(paragrafos_para_remover)} parágrafos vazios ou marcadores")
        
        return doc
    
    def _tipo_elemento(self, elemento):
        """
        Determina o tipo de elemento para logging detalhado.
        
        Args:
            elemento: Elemento do documento.
            
        Returns:
            String descritiva do elemento.
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        
        if isinstance(elemento, Paragraph):
            return f"Parágrafo: '{elemento.text[:30]}...'" if len(elemento.text) > 30 else f"Parágrafo: '{elemento.text}'"
        elif isinstance(elemento, Table):
            return f"Tabela com {len(elemento.rows)} linhas x {len(elemento.columns)} colunas"
        else:
            return f"Elemento tipo {type(elemento).__name__}"
    
    def _mapear_secoes_no_documento(self, doc: Document) -> Dict[str, Dict[str, Any]]:
        """
        Mapeia seções condicionais no documento.
        
        Procura por marcadores de seção no formato:
        {{#SECAO_ID}} para início de seção e {{/SECAO_ID}} para fim de seção.
        
        Args:
            doc: Documento a ser analisado.
            
        Returns:
            Dicionário com informações sobre cada seção encontrada.
        """
        mapeamento = {}
        secao_atual = None
        elemento_inicio = None
        elementos_secao_atual = []
        
        # Log detalhado para diagnóstico
        logger.debug("Iniciando mapeamento de seções no documento")
        
        # Regex para detectar início e fim de seção
        regex_inicio = r'{{[\s]*#([^{}]+?)[\s]*}}'
        regex_fim = r'{{[\s]*/([^{}]+?)[\s]*}}'
        
        secoes_abertas = {}  # Rastreia seções atualmente abertas
        
        # Contadores para diagnóstico
        total_paragrafos = len(doc.paragraphs)
        paragrafos_com_marcadores = 0
        marcadores_inicio_encontrados = 0
        marcadores_fim_encontrados = 0
        
        # Verifica todos os parágrafos para marcadores
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text
            
            # Log para parágrafos longos (amostra)
            if len(texto) > 100:
                texto_log = texto[:50] + "..." + texto[-50:]
            else:
                texto_log = texto
                
            logger.debug(f"Analisando parágrafo {i+1}/{total_paragrafos}: '{texto_log}'")
            
            # Busca marcadores no texto do parágrafo
            marcadores_inicio = re.finditer(regex_inicio, texto)
            marcadores_fim = re.finditer(regex_fim, texto)
            
            # Processa marcadores de início
            for match in marcadores_inicio:
                secao_id = match.group(1).strip()
                marcadores_inicio_encontrados += 1
                paragrafos_com_marcadores += 1
                
                logger.info(f"Marcador de INÍCIO de seção encontrado: '{secao_id}' (parágrafo {i+1})")
                
                # Registra início da seção
                secoes_abertas[secao_id] = {
                    'paragrafo_inicio': i,
                    'elemento_inicio': paragrafo,
                    'elementos': []
                }
            
            # Processa marcadores de fim
            for match in marcadores_fim:
                secao_id = match.group(1).strip()
                marcadores_fim_encontrados += 1
                
                # Apenas incrementa paragrafos_com_marcadores se não contou antes
                if not re.search(regex_inicio, texto):
                    paragrafos_com_marcadores += 1
                
                logger.info(f"Marcador de FIM de seção encontrado: '{secao_id}' (parágrafo {i+1})")
                
                # Verifica se a seção foi aberta
                if secao_id in secoes_abertas:
                    info_secao = secoes_abertas[secao_id]
                    
                    # Coleta elementos entre início e fim
                    elementos = []
                    for j in range(info_secao['paragrafo_inicio'], i + 1):
                        if j < len(doc.paragraphs):
                            elementos.append(doc.paragraphs[j])
                    
                    # Adiciona ao mapeamento final
                    mapeamento[secao_id] = {
                        'inicio': info_secao['elemento_inicio'],
                        'fim': paragrafo,
                        'elementos': elementos
                    }
                    
                    # Remove da lista de seções abertas
                    del secoes_abertas[secao_id]
                    
                    logger.info(f"Seção '{secao_id}' mapeada completamente com {len(elementos)} elementos")
                else:
                    logger.warning(f"Marcador de FIM para seção não aberta: '{secao_id}'")
        
        # Verifica se ficaram seções abertas não fechadas
        for secao_id, info in secoes_abertas.items():
            logger.warning(f"ALERTA: Seção '{secao_id}' foi aberta mas não fechada (parágrafo {info['paragrafo_inicio']+1})")
            
            # Mantém no mapeamento com os elementos até o fim do documento
            elementos = []
            for j in range(info['paragrafo_inicio'], len(doc.paragraphs)):
                elementos.append(doc.paragraphs[j])
                
            mapeamento[secao_id] = {
                'inicio': info['elemento_inicio'],
                'fim': doc.paragraphs[-1] if doc.paragraphs else None,
                'elementos': elementos,
                'incompleta': True  # Marca como incompleta para tratamento especial
            }
        
        # Log de resumo do mapeamento
        logger.info(f"Mapeamento de seções concluído:")
        logger.info(f"  Total de parágrafos analisados: {total_paragrafos}")
        logger.info(f"  Parágrafos com marcadores: {paragrafos_com_marcadores}")
        logger.info(f"  Marcadores de início encontrados: {marcadores_inicio_encontrados}")
        logger.info(f"  Marcadores de fim encontrados: {marcadores_fim_encontrados}")
        logger.info(f"  Seções mapeadas: {len(mapeamento)}")
        
        # Se não encontrou marcadores explícitos, tenta usar mapeamento por títulos
        if not mapeamento:
            logger.info("Nenhum marcador explícito encontrado. Tentando mapeamento por títulos...")
            mapeamento = self._mapear_secoes_por_titulos(doc, {})
            
        return mapeamento
        
    def _mapear_secoes_por_titulos(self, doc: Document, secoes_esperadas: Dict[str, str]) -> Dict[str, Dict[str, Any]]:
        """
        Mapeia seções no documento baseado em títulos ou headers.
        Usado como fallback quando marcadores explícitos não são encontrados.
        
        Args:
            doc: Documento a ser analisado.
            secoes_esperadas: Dicionário com mapeamento de ID de seção para campo condicional.
            
        Returns:
            Dicionário com informações sobre cada seção encontrada.
        """
        # Primeiro usamos o injetor de marcadores para identificar seções pelo conteúdo
        doc = self._injetar_marcadores_secao(doc)
        
        # Obtém as seções identificadas pelo método de injeção
        mapeamento = {}
        
        for secao_id, info in self.secoes_identificadas_pelo_conteudo.items():
            inicio_elemento = doc.paragraphs[info['titulo_idx']]
            
            # Determina elemento final
            fim_idx = info.get('fim_idx', len(doc.paragraphs) - 1)
            fim_elemento = doc.paragraphs[fim_idx] if fim_idx < len(doc.paragraphs) else None
            
            # Coleta elementos da seção
            elementos = []
            for i in range(info['titulo_idx'], fim_idx + 1):
                if i < len(doc.paragraphs):
                    elementos.append(doc.paragraphs[i])
            
            # Adiciona ao mapeamento
            mapeamento[secao_id] = {
                'inicio': inicio_elemento,
                'fim': fim_elemento,
                'elementos': elementos,
                'confianca': info.get('confianca', 'média')
            }
            
            # Diagnóstico
            logger.info(f"Seção '{secao_id}' mapeada com {len(elementos)} elementos")
            
            # Verifica se a seção está nas seções ativas
            if secao_id in self.secoes_ativas:
                logger.info(f"  → Seção '{secao_id}' está ATIVA")
            else:
                logger.info(f"  → Seção '{secao_id}' está INATIVA")
                
                # Diagnóstico adicional para entender por que está inativa
                if secao_id in self.mapeamento_secoes:
                    info_mapeamento = self.mapeamento_secoes[secao_id]
                    campo_condicional = info_mapeamento['campo_condicional']
                    valor_ativo = info_mapeamento['valor_ativo']
                    
                    if campo_condicional in secoes_esperadas:
                        logger.info(f"    Campo condicional '{campo_condicional}' está presente nos dados")
                        logger.info(f"    Valor esperado: '{valor_ativo}'")
                    else:
                        logger.warning(f"    Campo condicional '{campo_condicional}' NÃO está presente nos dados")
        
        return mapeamento
    
    def _processar_paragrafo(self, paragrafo, dados: Dict[str, Any]) -> None:
        """
        Processa um parágrafo, substituindo os campos pelos valores correspondentes.
        
        Args:
            paragrafo: Parágrafo a ser processado.
            dados: Dicionário com os dados para substituição.
        """
        # DIAGNÓSTICO: Log hexadecimal para detectar caracteres invisíveis
        texto_hex = paragrafo.text.encode('utf-8').hex()
        if paragrafo.text and len(paragrafo.text) < 100:  # Limita logs para textos menores
            logger.debug(f"Texto HEX: {texto_hex}")
        
        # ESTRATÉGIA 1: Tentativa convencional pelo texto completo do parágrafo
        texto_original = paragrafo.text
        
        # MELHORADO: Usar regex mais flexível para capturar placeholders com variações
        # Permite espaços entre chaves e nome do campo
        campos_no_paragrafo = re.findall(r'{{[\s]*([^{}]+?)[\s]*}}', texto_original)
        
        # Log de diagnóstico para verificar campos encontrados
        if campos_no_paragrafo:
            logger.debug(f"Campos encontrados no parágrafo: {campos_no_paragrafo}")
        
        # Registra campos encontrados que não são marcadores de seção
        for campo in campos_no_paragrafo:
            campo_limpo = campo.strip()
            # Ignora marcadores de seção (que começam com # ou /)
            if not campo_limpo.startswith('#') and not campo_limpo.startswith('/'):
                self.campos_encontrados.add(campo_limpo)
                logger.debug(f"Campo registrado: '{campo_limpo}'")
        
        # Se não há campos, tenta verificar runs fragmentados
        if not campos_no_paragrafo and len(paragrafo.runs) > 1:
            logger.debug(f"Parágrafo sem campos. Verificando runs fragmentados ({len(paragrafo.runs)} runs)")
            self._processar_runs_fragmentados(paragrafo, dados)
            return
        
        # Se não há campos, não precisa processar
        if not campos_no_paragrafo:
            return
        
        # Substitui os campos no texto completo
        texto_substituido = self._substituir_campos(texto_original, dados)
        
        # Se o texto foi alterado, atualiza o parágrafo preservando formatação
        if texto_substituido != texto_original:
            logger.debug(f"Texto substituído: '{texto_substituido}'")
            self._substituir_texto_preservando_formatacao(paragrafo, texto_original, texto_substituido)
    
    def _substituir_texto_preservando_formatacao(self, paragrafo, texto_original: str, texto_substituido: str) -> None:
        """
        Substitui o texto de um parágrafo preservando a formatação.
        
        Args:
            paragrafo: Parágrafo a ser modificado.
            texto_original: Texto original do parágrafo.
            texto_substituido: Texto com campos substituídos.
        """
        # O problema está aqui: estamos removendo todo o texto original
        # e substituindo apenas pelos valores, perdendo o texto entre placeholders
        
        # Se o parágrafo tem apenas uma run, é mais simples
        if len(paragrafo.runs) == 1:
            paragrafo.runs[0].text = texto_substituido
            return
        
        # CORREÇÃO: Em vez de limpar todas as runs, vamos apenas atualizar a primeira
        # com o texto completo substituído e depois limpar as demais
        if len(paragrafo.runs) > 0:
            # Atualiza a primeira run com o texto completo substituído
            paragrafo.runs[0].text = texto_substituido
            
            # Limpa todas as outras runs para evitar duplicação de texto
            for i in range(1, len(paragrafo.runs)):
                paragrafo.runs[i].text = ""
    
    def _processar_runs_fragmentados(self, paragrafo, dados: Dict[str, Any]) -> bool:
        """
        Processa placeholders que podem estar fragmentados entre diferentes runs.
        
        Args:
            paragrafo: Parágrafo a ser processado.
            dados: Dicionário com os dados para substituição.
            
        Returns:
            True se algum placeholder foi processado, False caso contrário.
        """
        # Log detalhado das runs para diagnóstico
        logger.debug(f"Verificando fragmentação em parágrafo com {len(paragrafo.runs)} runs")
        for i, run in enumerate(paragrafo.runs):
            texto_run_hex = run.text.encode('utf-8').hex()
            logger.debug(f"Run {i}: '{run.text}' (Hex: {texto_run_hex})")
        
        # 1. Reconstruir o texto concatenado das runs
        runs_texto = [run.text for run in paragrafo.runs]
        texto_completo = "".join(runs_texto)
        
        # Log do texto completo para diagnóstico
        texto_completo_hex = texto_completo.encode('utf-8').hex()
        logger.debug(f"Texto completo concatenado: '{texto_completo}' (Hex: {texto_completo_hex})")
        
        # 2. Buscar placeholders no texto completo com regex mais flexível
        # ALTERADO: Permite espaços entre chaves e nome do campo
        placeholders = re.finditer(r'{{[\s]*([^{}]+?)[\s]*}}', texto_completo)
        
        # Flag para indicar se algum placeholder foi encontrado e processado
        processou_algum = False
        
        # 3. Para cada placeholder, determinar quais runs são afetadas
        for match in placeholders:
            inicio = match.start()  # Início do placeholder (posição do primeiro '{')
            fim = match.end()       # Fim do placeholder (posição após o último '}')
            nome_campo = match.group(1).strip()
            placeholder_completo = match.group(0)
            
            logger.debug(f"Placeholder fragmentado encontrado: '{placeholder_completo}' (campo: '{nome_campo}')")
            logger.debug(f"Posição: {inicio}-{fim} (comprimento texto: {len(texto_completo)})")
            
            # Ignorar marcadores de seção
            if nome_campo.startswith('#') or nome_campo.startswith('/'):
                logger.debug(f"Ignorando marcador de seção: '{nome_campo}'")
                continue
            
            # Registra campo encontrado
            self.campos_encontrados.add(nome_campo)
            processou_algum = True
            
            # Determinar em quais runs o placeholder está
            runs_afetadas = []
            indices_em_runs = []
            posicao_atual = 0
            
            # Log para diagnóstico antes de processar runs
            logger.debug(f"Analisando distribuição do placeholder entre runs:")
            
            for i_run, texto_run in enumerate(runs_texto):
                nova_posicao = posicao_atual + len(texto_run)
                logger.debug(f"  Run {i_run}: posição {posicao_atual}-{nova_posicao}")
                
                # Run contém início do placeholder
                if posicao_atual <= inicio < nova_posicao:
                    runs_afetadas.append(i_run)
                    indices_em_runs.append((inicio - posicao_atual, min(len(texto_run), fim - posicao_atual)))
                    logger.debug(f"    Contém INÍCIO do placeholder em {inicio - posicao_atual}")
                
                # Run contém parte intermediária do placeholder
                elif posicao_atual > inicio and nova_posicao < fim:
                    runs_afetadas.append(i_run)
                    indices_em_runs.append((0, len(texto_run)))
                    logger.debug(f"    Contém PARTE INTERMEDIÁRIA do placeholder")
                
                # Run contém final do placeholder
                elif posicao_atual < fim <= nova_posicao and posicao_atual > inicio:
                    runs_afetadas.append(i_run)
                    indices_em_runs.append((0, fim - posicao_atual))
                    logger.debug(f"    Contém FIM do placeholder em {fim - posicao_atual}")
                
                posicao_atual = nova_posicao
            
            logger.debug(f"Runs afetadas: {runs_afetadas}")
            logger.debug(f"Índices em runs: {indices_em_runs}")
            
            # 4. Substituir o placeholder nas runs afetadas
            if nome_campo in dados:
                # Marcar como substituído
                self.campos_substituidos.add(nome_campo)
                
                # Obter valor para substituição
                valor = dados[nome_campo]
                
                # MELHORIA: Detecção aprimorada de campos monetários
                # Busca informações do campo para aplicar formatação
                campo_info = self.motor_regras.obter_campo_por_nome(nome_campo)
                
                # Tenta usar tipo de dado do modelo relacional
                tipo_monetario = False
                tipo_formatacao = None
                
                if campo_info:
                    # Verificar tipo_formatacao explícito
                    if 'tipo_formatacao' in campo_info:
                        tipo_formatacao = campo_info['tipo_formatacao']
                        # Se tem formato monetário (#.##0,00)
                        if tipo_formatacao and ('#.##0,00' in tipo_formatacao or 'dinheiro' in tipo_formatacao.lower() or 'monetário' in tipo_formatacao.lower()):
                            tipo_monetario = True
                    
                    # Verificar tipo_dado_programacao
                    if 'tipo_dado_programacao' in campo_info:
                        tipo = campo_info['tipo_dado_programacao'].lower()
                        tipo_monetario = tipo_monetario or 'dinheiro' in tipo or 'moeda' in tipo or 'valor' in tipo or 'salario' in tipo
                
                # Se não identificou pelo tipo, tenta identificar pelo nome do campo
                if not tipo_monetario:
                    nomes_monetarios = ['valor', 'salario', 'remuneracao', 'vencimento', 'subsidio', 'proventos']
                    tipo_monetario = any(termo in nome_campo.lower() for termo in nomes_monetarios)
                
                # Formatação para valores monetários
                valor_formatado = str(valor) if valor is not None else ""
                
                if tipo_monetario and isinstance(valor, (int, float)):
                    valor_formatado = self._formatar_valor_monetario(valor, tipo_formatacao)
                    logger.debug(f"Valor monetário formatado: '{valor}' -> '{valor_formatado}'")
                elif campo_info and 'tipo_dado_programacao' in campo_info:
                    tipo = campo_info['tipo_dado_programacao'].lower()
                    # Formatação para valores por extenso
                    if 'extenso' in tipo and isinstance(valor, (int, float)):
                        valor_formatado = self._valor_por_extenso(valor)
                
                # Aplicar substituição nas runs
                if runs_afetadas:
                    # Primeira run com parte do placeholder
                    primeira_run = paragrafo.runs[runs_afetadas[0]]
                    primeiro_inicio, primeiro_fim = indices_em_runs[0]
                    
                    logger.debug(f"Substituindo em {len(runs_afetadas)} runs por '{valor_formatado}'")
                    
                    # A primeira run mantém o texto antes do placeholder e recebe o valor
                    novo_texto_primeira_run = primeira_run.text[:primeiro_inicio] + valor_formatado
                    logger.debug(f"Primeira run original: '{primeira_run.text}'")
                    logger.debug(f"Primeira run nova: '{novo_texto_primeira_run}'")
                    primeira_run.text = novo_texto_primeira_run
                    
                    # As runs intermediárias são limpas
                    for i in range(1, len(runs_afetadas)):
                        run_idx = runs_afetadas[i]
                        logger.debug(f"Limpando run {run_idx}: '{paragrafo.runs[run_idx].text}'")
                        paragrafo.runs[run_idx].text = ""
                        
                    # A última run mantém o texto após o placeholder
                    if len(runs_afetadas) > 1:
                        ultima_run = paragrafo.runs[runs_afetadas[-1]]
                        _, ultimo_fim = indices_em_runs[-1]
                        
                        if ultimo_fim < len(ultima_run.text):
                            novo_texto_ultima_run = ultima_run.text[ultimo_fim:]
                            logger.debug(f"Última run original: '{ultima_run.text}'")
                            logger.debug(f"Última run nova: '{novo_texto_ultima_run}'")
                            ultima_run.text = novo_texto_ultima_run
                        else:
                            logger.debug(f"Última run limpada completamente: '{ultima_run.text}'")
                            ultima_run.text = ""
                    
                    logger.info(f"Substituído placeholder fragmentado '{nome_campo}' por '{valor_formatado}'")
            else:
                # Campo ausente
                self.campos_ausentes.add(nome_campo)
                
                # Verifica se é obrigatório
                campo_info = self.motor_regras.obter_campo_por_nome(nome_campo)
                obrigatorio = False
                if campo_info and 'obrigatorio_quando_ativo' in campo_info:
                    obrigatorio = campo_info['obrigatorio_quando_ativo']
                
                # Substitui por texto indicando campo obrigatório ausente ou vazio
                texto_substituicao = f"**[CAMPO OBRIGATÓRIO: {nome_campo}]**" if obrigatorio else ""
                
                if obrigatorio:
                    self.campos_obrigatorios_ausentes.add(nome_campo)
                    logger.warning(f"Campo obrigatório ausente: '{nome_campo}'")
                else:
                    logger.debug(f"Campo ausente (não obrigatório): '{nome_campo}'")
                
                # Similar à substituição normal, mas com texto diferente
                if runs_afetadas:
                    primeira_run = paragrafo.runs[runs_afetadas[0]]
                    primeiro_inicio, primeiro_fim = indices_em_runs[0]
                    
                    logger.debug(f"Substituindo campo ausente em {len(runs_afetadas)} runs")
                    primeira_run.text = primeira_run.text[:primeiro_inicio] + texto_substituicao
                    
                    for i in range(1, len(runs_afetadas)):
                        paragrafo.runs[runs_afetadas[i]].text = ""
                        
                    if len(runs_afetadas) > 1:
                        ultima_run = paragrafo.runs[runs_afetadas[-1]]
                        _, ultimo_fim = indices_em_runs[-1]
                        if ultimo_fim < len(ultima_run.text):
                            ultima_run.text = ultima_run.text[ultimo_fim:]
                        else:
                            ultima_run.text = ""
        
        return processou_algum
    
    def _substituir_campos(self, texto: str, dados: Dict[str, Any]) -> str:
        """
        Substitui os marcadores de campo no texto pelos valores correspondentes.
        
        Args:
            texto: Texto com marcadores.
            dados: Dicionário com os dados para substituição.
            
        Returns:
            Texto com os marcadores substituídos.
        """
        def substituir(match):
            # Captura o placeholder completo e o nome do campo
            placeholder_completo = match.group(0)
            nome_campo = match.group(1).strip()  # Remove espaços extras
            
            logger.debug(f"Substituindo placeholder: '{placeholder_completo}', campo: '{nome_campo}'")
            
            # Ignora marcadores de seção
            if nome_campo.startswith('#') or nome_campo.startswith('/'):
                logger.debug(f"Ignorando marcador de seção: '{nome_campo}'")
                return match.group(0)
            
            # Busca informações do campo no modelo relacional    
            campo_info = self.motor_regras.obter_campo_por_nome(nome_campo)
            
            if nome_campo in dados:
                valor = dados[nome_campo]
                self.campos_substituidos.add(nome_campo)
                
                logger.debug(f"Valor encontrado para '{nome_campo}': '{valor}'")
                
                # DETECÇÃO DE CAMPO MONETÁRIO (APRIMORADA)
                tipo_monetario = False
                tipo_formatacao = None
                
                # 1. Verificar por tipo_formatacao explícito no campo_info
                if campo_info and 'tipo_formatacao' in campo_info:
                    tipo_formatacao = campo_info['tipo_formatacao']
                    if tipo_formatacao:
                        indicadores_moeda = ['#.##0,00', 'dinheiro', 'monetário', 'moeda', 'r$']
                        tipo_monetario = any(ind in tipo_formatacao.lower() for ind in indicadores_moeda)
                    
                # 2. Verificar por tipo_dado_programacao no campo_info
                if campo_info and 'tipo_dado_programacao' in campo_info and not tipo_monetario:
                    tipo = campo_info['tipo_dado_programacao'].lower()
                    indicadores_moeda = ['dinheiro', 'moeda', 'valor', 'salario', 'monetário']
                    tipo_monetario = any(ind in tipo for ind in indicadores_moeda)
                
                # 3. Identificação pelo nome do campo
                if not tipo_monetario:
                    nome_lower = nome_campo.lower()
                    nomes_monetarios = ['valor', 'salario', 'remuneracao', 'vencimento', 'subsidio', 'proventos', 
                                        'montante', 'importancia', 'quantia', 'bruto', 'liquido', 'total']
                    tipo_monetario = any(termo in nome_lower for termo in nomes_monetarios)
                
                # APLICAÇÃO DE FORMATAÇÃO MONETÁRIA
                if tipo_monetario and isinstance(valor, (int, float)):
                    texto_formatado = self._formatar_valor_monetario(valor, tipo_formatacao)
                    logger.debug(f"Formatado como moeda: '{valor}' → '{texto_formatado}'")
                    return texto_formatado
                
                # VALOR POR EXTENSO
                if campo_info and 'tipo_dado_programacao' in campo_info:
                    tipo = campo_info['tipo_dado_programacao'].lower()
                    if 'extenso' in tipo and isinstance(valor, (int, float)):
                        texto_formatado = self._valor_por_extenso(valor)
                        logger.debug(f"Formatado por extenso: '{valor}' → '{texto_formatado}'")
                        return texto_formatado
                
                # Retorna o valor como string
                return str(valor) if valor is not None else ""
            else:
                self.campos_ausentes.add(nome_campo)
                logger.debug(f"Campo ausente: '{nome_campo}'")
                
                # Verifica se o campo é obrigatório usando o modelo relacional
                obrigatorio = False
                if campo_info and 'obrigatorio_quando_ativo' in campo_info:
                    obrigatorio = campo_info['obrigatorio_quando_ativo']
                
                if obrigatorio:
                    self.campos_obrigatorios_ausentes.add(nome_campo)
                    # Marca visivelmente os campos obrigatórios ausentes
                    texto_marcacao = f"**[CAMPO OBRIGATÓRIO: {nome_campo}]**"
                    logger.warning(f"Campo obrigatório ausente: '{nome_campo}'")
                    return texto_marcacao
                else:
                    # Campos não obrigatórios ficam em branco
                    return ""
        
        # MELHORADO: Regex mais flexível que permite espaços entre chaves e o nome do campo
        regex_substituicao = r'{{[\s]*([^{}]+?)[\s]*}}'
        resultado = re.sub(regex_substituicao, substituir, texto)
        return resultado
    
    def _exibir_campos_ausentes(self) -> None:
        """
        Exibe informações sobre campos ausentes nos dados.
        """
        if not self.campos_ausentes:
            logger.info("Todos os campos foram substituídos com sucesso!")
            return
        
        # Organiza campos ausentes por categoria
        campos_por_categoria = self._agrupar_campos_por_categoria(self.campos_ausentes)
        
        # Log de campos ausentes organizados por categoria
        logger.warning(f"Total de campos ausentes: {len(self.campos_ausentes)} de {len(self.campos_encontrados)}")
        
        if self.campos_obrigatorios_ausentes:
            logger.error(f"CAMPOS OBRIGATÓRIOS AUSENTES: {len(self.campos_obrigatorios_ausentes)}")
            for campo in sorted(self.campos_obrigatorios_ausentes):
                logger.error(f"  - {campo}")
        
        for categoria, campos in campos_por_categoria.items():
            logger.warning(f"Categoria '{categoria}': {len(campos)} campos ausentes")
            for campo in campos:
                logger.warning(f"  - {campo}")
    
    def _agrupar_campos_por_categoria(self, campos: Set[str]) -> Dict[str, List[str]]:
        """
        Agrupa campos por categoria usando o modelo relacional.
        
        Args:
            campos: Conjunto de nomes de campos.
            
        Returns:
            Dicionário com campos agrupados por categoria.
        """
        # Categorias padrão
        categorias_padrao = {
            "Informações Processuais": [],
            "Partes": [],
            "Datas": [],
            "Valores": [],
            "Cálculos": [],
            "Pedidos": [],
            "Documentos": []
        }
        categoria_desconhecida = "Outros"
        
        # Inicializa o resultado com as categorias padrão
        resultado = categorias_padrao.copy()
        resultado[categoria_desconhecida] = []
        
        # Prefixos conhecidos para categorização automática quando não há categoria no modelo relacional
        prefixos_categorias = {
            "proc": "Informações Processuais",
            "autor": "Partes",
            "reu": "Partes",
            "reclamante": "Partes",
            "reclamado": "Partes",
            "data": "Datas",
            "valor": "Valores",
            "calculo": "Cálculos",
            "pedido": "Pedidos",
            "documento": "Documentos"
        }
        
        for campo in campos:
            # Tenta obter informações do campo usando o motor de regras
            campo_info = self.motor_regras.obter_campo_por_nome(campo)
            
            # Se encontrou informações e tem categoria
            if campo_info and 'categoria' in campo_info and campo_info['categoria']:
                categoria = campo_info['categoria']
                # Verifica se a categoria já existe no resultado, senão cria
                if categoria not in resultado:
                    resultado[categoria] = []
                resultado[categoria].append(campo)
                continue
            
            # Se não achou no modelo relacional, tenta categorizar por prefixo
            categorizado = False
            campo_lower = campo.lower()
            for prefixo, categoria in prefixos_categorias.items():
                if campo_lower.startswith(prefixo):
                    resultado[categoria].append(campo)
                    categorizado = True
                    break
            
            # Se não conseguiu categorizar, coloca em Outros
            if not categorizado:
                resultado[categoria_desconhecida].append(campo)
        
        # Remove categorias vazias e ordena os campos em cada categoria
        categorias_final = {}
        for categoria, lista_campos in resultado.items():
            if lista_campos:  # Só mantém categorias não vazias
                categorias_final[categoria] = sorted(lista_campos)
        
        return categorias_final
    
    def _registrar_estatisticas(self, dados: Dict[str, Any]) -> None:
        """
        Registra estatísticas do processamento.
        
        Args:
            dados: Dicionário com os dados usados na substituição.
        """
        total_encontrados = len(self.campos_encontrados)
        total_substituidos = len(self.campos_substituidos)
        total_ausentes = len(self.campos_ausentes)
        total_obrigatorios_ausentes = len(self.campos_obrigatorios_ausentes)
        total_secoes_encontradas = len(self.secoes_encontradas)
        total_secoes_ativas = len(self.secoes_ativas)
        
        # Calcula a porcentagem de completude
        porcentagem_completude = 0
        if total_encontrados > 0:
            porcentagem_completude = (total_substituidos / total_encontrados) * 100
        
        # Determina o status do processamento
        status = "Erro"
        if total_obrigatorios_ausentes == 0:
            if porcentagem_completude >= 95:
                status = "Sucesso"
            elif porcentagem_completude >= 50:
                status = "Parcial"
        
        # Organiza campos ausentes por categoria
        campos_por_categoria = self._agrupar_campos_por_categoria(self.campos_ausentes)
        
        # Cria dicionário de estatísticas
        self.estatisticas_processamento = {
            "status": status,
            "total_campos_encontrados": total_encontrados,
            "total_campos_substituidos": total_substituidos,
            "total_campos_ausentes": total_ausentes,
            "total_campos_obrigatorios_ausentes": total_obrigatorios_ausentes,
            "campos_ausentes_por_categoria": campos_por_categoria,
            "porcentagem_completude": porcentagem_completude,
            "total_campos_dados": len(dados),
            "total_secoes_encontradas": total_secoes_encontradas,
            "total_secoes_ativas": total_secoes_ativas
        }
        
        # Registra no log
        logger.info(f"Estatísticas de processamento:")
        logger.info(f"  - Status: {status}")
        logger.info(f"  - Completude: {porcentagem_completude:.2f}%")
        logger.info(f"  - Campos encontrados: {total_encontrados}")
        logger.info(f"  - Campos substituídos: {total_substituidos}")
        logger.info(f"  - Campos ausentes: {total_ausentes}")
        logger.info(f"  - Campos obrigatórios ausentes: {total_obrigatorios_ausentes}")
        
        # Log detalhado por categoria
        if total_ausentes > 0:
            logger.info(f"Campos ausentes por categoria:")
            for categoria, campos in campos_por_categoria.items():
                logger.info(f"  - {categoria}: {len(campos)} campos")
                for campo in sorted(campos)[:10]:  # Mostra apenas os 10 primeiros para não poluir o log
                    logger.info(f"      * {campo}")
                if len(campos) > 10:
                    logger.info(f"      * ... e mais {len(campos) - 10} campos")
        
        logger.info(f"  - Seções encontradas: {total_secoes_encontradas}")
        logger.info(f"  - Seções ativas: {total_secoes_ativas}")
    
    def obter_estatisticas(self) -> Dict[str, Any]:
        """
        Obtém as estatísticas do último processamento.
        
        Returns:
            Dicionário com estatísticas.
        """
        return self.estatisticas_processamento
    
    def _valor_por_extenso(self, valor: float) -> str:
        """
        Converte um valor numérico para sua representação por extenso.
        Implementação mais completa.
        
        Args:
            valor: Valor numérico.
            
        Returns:
            String com o valor por extenso.
        """
        unidades = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
        dezenas = ["", "dez", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
        centenas = ["", "cem", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
        especiais = {
            11: "onze", 12: "doze", 13: "treze", 14: "quatorze", 15: "quinze",
            16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove"
        }
        
        # Separa parte inteira e decimal
        parte_inteira = int(valor)
        parte_decimal = int(round((valor - parte_inteira) * 100))
        
        # Função para processar um número até 999
        def processar_grupo(num):
            if num == 0:
                return ""
            if num <= 9:
                return unidades[num]
            if num in especiais:
                return especiais[num]
            if num < 100:
                dezena = dezenas[num // 10]
                unidade = unidades[num % 10]
                if unidade:
                    return f"{dezena} e {unidade}"
                return dezena
            # Números de 100 a 999
            if num == 100:
                return "cem"
            centena = centenas[num // 100]
            resto = num % 100
            if resto == 0:
                return centena
            return f"{centena} e {processar_grupo(resto)}"
        
        # Processa parte inteira
        if parte_inteira == 0:
            texto_reais = "zero"
        else:
            texto_reais = processar_grupo(parte_inteira)
        
        # Processa parte decimal
        if parte_decimal == 0:
            texto_centavos = "zero"
        else:
            texto_centavos = processar_grupo(parte_decimal)
        
        # Monta texto final
        if parte_inteira == 1:
            texto_reais += " real"
        else:
            texto_reais += " reais"
            
        if parte_decimal == 1:
            texto_centavos += " centavo"
        else:
            texto_centavos += " centavos"
            
        return f"{texto_reais} e {texto_centavos}"

    def _injetar_marcadores_secao(self, doc: Document) -> Document:
        """
        Injeta marcadores de seção no documento para processamento condicional.
        Usado quando o documento não tem marcadores explícitos.
        
        Args:
            doc: Documento a ser processado.
            
        Returns:
            Documento com marcadores injetados.
        """
        logger.info("Iniciando injeção de marcadores de seção no documento")
        
        # Busca parágrafos que podem ser títulos de seção
        secoes_identificadas = {}
        pontuacoes_secoes = {}  # Para registrar a confiança da identificação
        
        # Análise refinada para identificação de títulos
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text.strip()
            
            # Pular parágrafos vazios
            if not texto:
                continue
                
            # Calcular características que indicam um título
            pontuacao = 0
            caracteristicas = []
            
            # Verificar se está em maiúsculas
            if texto.isupper() and len(texto) > 3:
                pontuacao += 3
                caracteristicas.append("MAIÚSCULAS")
            
            # Verificar formato de texto (negrito, etc.)
            tem_negrito = False
            tem_italico = False
            for run in paragrafo.runs:
                if run.bold:
                    tem_negrito = True
                if run.italic:
                    tem_italico = True
            
            if tem_negrito:
                pontuacao += 3
                caracteristicas.append("NEGRITO")
            if tem_italico:
                pontuacao += 1
                caracteristicas.append("ITÁLICO")
                
            # Verificar estilo do parágrafo
            if paragrafo.style and paragrafo.style.name:
                estilo = paragrafo.style.name.lower()
                if 'heading' in estilo or 'título' in estilo or 'title' in estilo:
                    pontuacao += 5
                    caracteristicas.append(f"ESTILO={paragrafo.style.name}")
            
            # Comprimento do texto (títulos geralmente são curtos)
            if len(texto) < 50:
                pontuacao += 1
            
            # Verificar se começa com numeral (comum em títulos jurídicos)
            if re.match(r'^\d+[\.\)\-]', texto):
                pontuacao += 2
                caracteristicas.append("NUMERAÇÃO")
            
            # Verificar se este parágrafo corresponde a alguma seção do mapeamento
            for secao_id, info in self.mapeamento_secoes.items():
                # Pontuação base por correspondência com palavras-chave
                pontuacao_secao = 0
                
                # Verificar correspondência com palavras-chave
                for keyword in info['palavras_chave']:
                    if keyword.lower() in texto.lower():
                        pontuacao_secao += 3
                        caracteristicas.append(f"KEYWORD={keyword}")
                
                # Se tiver pontuação suficiente (mais de uma palavra-chave)
                if pontuacao_secao >= 3:
                    # Combina a pontuação do formato (título) com a da correspondência de conteúdo
                    pontuacao_final = pontuacao + pontuacao_secao
                    
                    # Registra se for uma correspondência boa
                    if pontuacao_final >= 5:
                        # Verifica se já encontrou esta seção antes com menor pontuação
                        if secao_id in pontuacoes_secoes and pontuacoes_secoes[secao_id] >= pontuacao_final:
                            continue
                            
                        # Registra ou atualiza a seção
                        secoes_identificadas[secao_id] = {
                            'titulo_idx': i,
                            'titulo_texto': texto,
                            'caracteristicas': ', '.join(caracteristicas)
                        }
                        pontuacoes_secoes[secao_id] = pontuacao_final
                        
                        logger.info(f"Seção '{secao_id}' identificada no parágrafo {i+1}: '{texto}'")
                        logger.info(f"  Pontuação: {pontuacao_final} | Características: {', '.join(caracteristicas)}")
        
        # Não faz modificações se não encontrou seções
        if not secoes_identificadas:
            logger.warning("Nenhuma seção identificada para injeção de marcadores")
            return doc
            
        # Organiza as seções por ordem de aparecimento no documento
        secoes_ordenadas = sorted(secoes_identificadas.items(), key=lambda x: x[1]['titulo_idx'])
        
        # Registra a ordem identificada
        logger.info(f"Ordem de seções identificadas no documento:")
        for i, (secao_id, info) in enumerate(secoes_ordenadas):
            logger.info(f"  {i+1}. {secao_id} (parágrafo {info['titulo_idx']+1}): '{info['titulo_texto']}'")
        
        # Para cada seção identificada, precisamos determinar seu fim
        for i, (secao_id, info) in enumerate(secoes_ordenadas):
            inicio_idx = info['titulo_idx']
            
            # Determina o fim da seção (próximo título ou fim do documento)
            if i < len(secoes_ordenadas) - 1:
                fim_idx = secoes_ordenadas[i+1][1]['titulo_idx'] - 1
            else:
                fim_idx = len(doc.paragraphs) - 1
                
            # Registra os índices
            secoes_identificadas[secao_id]['fim_idx'] = fim_idx
            
            # Tamanho da seção em parágrafos
            tamanho_secao = fim_idx - inicio_idx + 1
            secoes_identificadas[secao_id]['tamanho'] = tamanho_secao
            
            logger.info(f"Seção '{secao_id}': parágrafos {inicio_idx+1} a {fim_idx+1} ({tamanho_secao} parágrafos)")
            
            # Amostra do conteúdo para diagnóstico
            amostra_texto = []
            for j in range(inicio_idx, min(inicio_idx + 3, fim_idx + 1)):
                if j < len(doc.paragraphs):
                    texto_amostra = doc.paragraphs[j].text.strip()
                    if texto_amostra:
                        amostra_texto.append(f"P{j+1}: '{texto_amostra[:50]}...'")
            
            if amostra_texto:
                logger.debug(f"Amostra da seção '{secao_id}':\n  " + "\n  ".join(amostra_texto))
            
        # IMPORTANTE: Não modificamos o documento diretamente aqui
        # Em vez disso, registramos os limites das seções para usar durante o processamento
        self.secoes_identificadas_pelo_conteudo = secoes_identificadas
        
        return doc

    def _formatar_valor_monetario(self, valor, tipo_formatacao=None):
        """
        Formata um valor monetário de acordo com o padrão especificado.
        
        Args:
            valor: Valor numérico a ser formatado
            tipo_formatacao: Padrão de formatação (#.##0,00 para BRL)
            
        Returns:
            String formatada do valor monetário
        """
        if not isinstance(valor, (int, float)):
            return str(valor)
            
        # Import locale para formatação conforme configuração regional
        import locale
        try:
            # Tenta configurar o locale para pt_BR (português do Brasil)
            locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
        except locale.Error:
            try:
                # Fallback para Portuguese/Brazil
                locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')
            except locale.Error:
                # Se não conseguir, mantém o locale atual
                pass
                
        # Estratégias de formatação em ordem de preferência
        try:
            # Tentativa 1: Usar formato brasileiro com separador de milhar via locale
            texto_formatado = f"R$ {locale.currency(valor, grouping=True, symbol=False)}"
        except:
            try:
                # Tentativa 2: Formato manual com separador de milhar
                texto_formatado = f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            except:
                # Tentativa 3: Formato simples
                texto_formatado = f"R$ {valor:.2f}".replace('.', ',')
            
        logger.debug(f"Valor monetário formatado: {valor} → {texto_formatado} (formato: {tipo_formatacao})")
        return texto_formatado

    def _determinar_secoes_ativas(self, dados: Dict[str, Any]) -> List[str]:
        """
        Determina quais seções condicionais devem ser ativas baseado nos dados.
        
        Args:
            dados: Dicionário com os dados de contexto.
            
        Returns:
            Lista de IDs de seções que devem estar ativas.
        """
        secoes_ativas = []
        
        # Logs detalhados para diagnóstico
        logger.info("Determinando seções ativas com base nos dados disponíveis")
        logger.debug(f"Dados disponíveis: {list(dados.keys())}")
        
        # Verifica cada seção mapeada
        for secao_id, info in self.mapeamento_secoes.items():
            campo_condicional = info['campo_condicional']
            valor_ativo = info['valor_ativo']
            descricao = info.get('descricao', secao_id)
            
            # Verifica se o campo condicional existe nos dados
            if campo_condicional in dados:
                valor_real = dados[campo_condicional]
                logger.debug(f"Seção '{secao_id}' ({descricao}): campo '{campo_condicional}' = '{valor_real}'")
                
                # Verificação dos casos típicos (Sim/Não, presença/ausência)
                if valor_ativo == 'Sim' and valor_real in ['Sim', 'sim', 'S', 's', True, 1, '1']:
                    secoes_ativas.append(secao_id)
                    logger.info(f"Seção '{secao_id}' ({descricao}) ATIVADA: '{campo_condicional}' = '{valor_real}'")
                elif valor_ativo == 'Não' and valor_real in ['Não', 'não', 'nao', 'N', 'n', False, 0, '0']:
                    secoes_ativas.append(secao_id)
                    logger.info(f"Seção '{secao_id}' ({descricao}) ATIVADA: '{campo_condicional}' = '{valor_real}'")
                # Outros casos de correspondência direta
                elif str(valor_real).lower() == str(valor_ativo).lower():
                    secoes_ativas.append(secao_id)
                    logger.info(f"Seção '{secao_id}' ({descricao}) ATIVADA: '{campo_condicional}' = '{valor_real}'")
                # Casos específicos conforme necessidade
                elif 'dias_aviso_previo' in campo_condicional and valor_real and str(valor_real) != '0':
                    secoes_ativas.append(secao_id)
                    logger.info(f"Seção '{secao_id}' ({descricao}) ATIVADA: '{campo_condicional}' = '{valor_real}'")
                else:
                    logger.info(f"Seção '{secao_id}' ({descricao}) DESATIVADA: '{campo_condicional}' = '{valor_real}' (esperado: '{valor_ativo}')")
            else:
                logger.warning(f"Seção '{secao_id}' ({descricao}) DESATIVADA: campo '{campo_condicional}' não encontrado nos dados")
        
        # Log final de diagnóstico
        if secoes_ativas:
            logger.info(f"Seções ativas determinadas: {secoes_ativas}")
        else:
            logger.warning("Nenhuma seção ativa determinada - documento será processado sem remoção de seções")
            
        return secoes_ativas 

    def _validar_documento(self, dados: Dict[str, Any]) -> None:
        """
        Realiza verificações finais no documento processado.
        
        Args:
            dados: Dicionário com os dados usados no processamento.
            
        Raises:
            DadosError: Se houver campos obrigatórios ausentes.
        """
        # Verifica se algum campo obrigatório está ausente
        if self.campos_obrigatorios_ausentes and self.modo_estrito:
            campos_faltantes = sorted(list(self.campos_obrigatorios_ausentes))
            mensagem = f"Campos obrigatórios ausentes: {', '.join(campos_faltantes)}"
            logger.error(mensagem)
            raise DadosError(mensagem)
            
        # Exibe informações sobre campos processados
        logger.info(f"Total de campos encontrados: {len(self.campos_encontrados)}")
        logger.info(f"Campos substituídos: {len(self.campos_substituidos)}")
        logger.info(f"Campos ausentes: {len(self.campos_ausentes)}")
        
        # Exibe detalhes sobre campos ausentes
        if self.campos_ausentes:
            campos_ausentes_lista = sorted(list(self.campos_ausentes))
            logger.warning(f"Os seguintes campos não foram encontrados nos dados: {', '.join(campos_ausentes_lista)}")
            
            # Verifica se os campos ausentes têm valores padrão no modelo relacional
            campos_com_padrao = []
            for campo in campos_ausentes_lista:
                info = self.motor_regras.obter_campo_por_nome(campo)
                if info and 'valor_padrao' in info and info['valor_padrao'] is not None:
                    campos_com_padrao.append(f"{campo} (padrão: {info['valor_padrao']})")
            
            if campos_com_padrao:
                logger.info(f"Os seguintes campos ausentes têm valores padrão: {', '.join(campos_com_padrao)}")
                
        # Exibe detalhes sobre campos obrigatórios ausentes
        if self.campos_obrigatorios_ausentes:
            campos_obrigatorios = sorted(list(self.campos_obrigatorios_ausentes))
            logger.warning(f"Os seguintes campos OBRIGATÓRIOS não foram fornecidos: {', '.join(campos_obrigatorios)}")
            
        # Exibe informações sobre seções
        logger.info(f"Seções encontradas no documento: {sorted(list(self.secoes_encontradas))}")
        logger.info(f"Seções ativas aplicadas: {sorted(self.secoes_ativas)}")
        
        # Verifica se há seções não reconhecidas
        secoes_nao_mapeadas = self.secoes_encontradas - set(self.mapeamento_secoes.keys())
        if secoes_nao_mapeadas:
            logger.warning(f"As seguintes seções foram encontradas, mas não estão no mapeamento: {sorted(list(secoes_nao_mapeadas))}")
            
        # Verifica se seções esperadas não foram encontradas no documento
        secoes_esperadas = set(self.mapeamento_secoes.keys())
        secoes_ausentes = secoes_esperadas - self.secoes_encontradas
        if secoes_ausentes:
            logger.warning(f"As seguintes seções estão no mapeamento, mas não foram encontradas no documento: {sorted(list(secoes_ausentes))}")
            
            # Verifica quais destas seções ausentes deveriam estar ativas
            secoes_ausentes_ativas = secoes_ausentes.intersection(set(self.secoes_ativas))
            if secoes_ausentes_ativas:
                logger.error(f"ATENÇÃO: As seguintes seções deveriam estar ativas, mas não foram encontradas no documento: {sorted(list(secoes_ausentes_ativas))}")
                
        # Verificação adicional para valores monetários
        campos_monetarios = []
        for campo in self.campos_substituidos:
            info = self.motor_regras.obter_campo_por_nome(campo)
            if info:
                # Verifica se é um campo monetário
                tipo_monetario = False
                if 'tipo_formatacao' in info:
                    tipo_formatacao = info['tipo_formatacao']
                    if tipo_formatacao and ('#.##0,00' in tipo_formatacao or 'dinheiro' in tipo_formatacao.lower() or 'monetário' in tipo_formatacao.lower()):
                        tipo_monetario = True
                
                if 'tipo_dado_programacao' in info:
                    tipo = info['tipo_dado_programacao'].lower()
                    if 'dinheiro' in tipo or 'moeda' in tipo or 'valor' in tipo or 'salario' in tipo:
                        tipo_monetario = True
                        
                if tipo_monetario:
                    campos_monetarios.append(campo)
                    
        if campos_monetarios:
            logger.info(f"Os seguintes campos monetários foram processados: {', '.join(campos_monetarios)}")
        
        return None 

    def _substituir_todos_campos(self, doc: Document, dados: Dict[str, Any]) -> Document:
        """
        Substitui todos os campos no documento pelos valores correspondentes.
        
        Args:
            doc: Documento a ser processado.
            dados: Dicionário com os dados para substituição.
            
        Returns:
            Documento com campos substituídos.
        """
        logger.info("Iniciando substituição de campos no documento")
        
        # Processa todos os parágrafos
        for i, paragrafo in enumerate(doc.paragraphs):
            # Primeiro verifica se há placeholders fragmentados
            processou_fragmentados = self._processar_runs_fragmentados(paragrafo, dados)
            
            # Se não processou fragmentados, processa o parágrafo inteiro
            if not processou_fragmentados:
                texto_original = paragrafo.text
                if '{{' in texto_original and '}}' in texto_original:
                    texto_substituido = self._substituir_campos(texto_original, dados)
                    
                    # Só aplica a substituição se houve mudança
                    if texto_substituido != texto_original:
                        paragrafo.text = texto_substituido
                        logger.debug(f"Parágrafo {i+1} substituído: '{texto_original[:50]}...' → '{texto_substituido[:50]}...'")
        
        # Processa tabelas
        for i, tabela in enumerate(doc.tables):
            logger.debug(f"Processando tabela {i+1} com {len(tabela.rows)} linhas")
            
            for row_idx, row in enumerate(tabela.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Processa cada parágrafo na célula
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        # Primeiro verifica se há placeholders fragmentados
                        processou_fragmentados = self._processar_runs_fragmentados(paragraph, dados)
                        
                        # Se não processou fragmentados, processa o parágrafo inteiro
                        if not processou_fragmentados:
                            texto_original = paragraph.text
                            if '{{' in texto_original and '}}' in texto_original:
                                texto_substituido = self._substituir_campos(texto_original, dados)
                                
                                # Só aplica a substituição se houve mudança
                                if texto_substituido != texto_original:
                                    paragraph.text = texto_substituido
                                    logger.debug(f"Tabela {i+1}, linha {row_idx+1}, coluna {cell_idx+1}, parágrafo {p_idx+1} substituído")
        
        # Processa cabeçalhos
        for section in doc.sections:
            # Cabeçalho primeiro
            if section.header:
                for paragraph in section.header.paragraphs:
                    processou_fragmentados = self._processar_runs_fragmentados(paragraph, dados)
                    if not processou_fragmentados:
                        texto_original = paragraph.text
                        if '{{' in texto_original and '}}' in texto_original:
                            texto_substituido = self._substituir_campos(texto_original, dados)
                            if texto_substituido != texto_original:
                                paragraph.text = texto_substituido
                                logger.debug(f"Cabeçalho substituído: '{texto_original[:50]}...' → '{texto_substituido[:50]}...'")
            
            # Rodapé depois
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    processou_fragmentados = self._processar_runs_fragmentados(paragraph, dados)
                    if not processou_fragmentados:
                        texto_original = paragraph.text
                        if '{{' in texto_original and '}}' in texto_original:
                            texto_substituido = self._substituir_campos(texto_original, dados)
                            if texto_substituido != texto_original:
                                paragraph.text = texto_substituido
                                logger.debug(f"Rodapé substituído: '{texto_original[:50]}...' → '{texto_substituido[:50]}...'")
        
        logger.info(f"Substituição de campos concluída. Encontrados {len(self.campos_encontrados)} campos, substituídos {len(self.campos_substituidos)}")
        return doc 