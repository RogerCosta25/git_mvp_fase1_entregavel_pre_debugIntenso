"""
Gerador de documentos Word para o sistema de peticionamento.
"""
import os
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Adiciona o diretório pai ao path para importar módulos
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config
from src.exceptions import (
    ArquivoNaoEncontradoError,
    TemplateError,
    SubstituicaoError
)
from src.logger import logger
from src.template_metadata import TemplateMetadata
from src.template_repository import TemplateRepository, FileSystemTemplateRepository

class GeradorDocumento:
    """
    Classe responsável pela geração de documentos Word com base em templates.
    """
    
    def __init__(self, caminho_template=None, repository: TemplateRepository=None):
        """
        Inicializa o gerador de documentos com repositório de templates.
        Args:
            caminho_template: Caminho para o template Word (opcional se repository for fornecido).
            repository: Instância de TemplateRepository (padrão FileSystemTemplateRepository).
        """
        # Configura repositório de templates
        if repository:
            self.repo = repository
        else:
            self.repo = FileSystemTemplateRepository(template_path=caminho_template)
        self.documento = None
        self.placeholders_encontrados = set()
        # Inicializa o metadata de placeholders
        self.metadata = TemplateMetadata()
    
    def carregar_template(self, caminho=None):
        """
        Carrega o template Word.
        
        Args:
            caminho: Caminho para o template. Se None, usa self.caminho_template.
            
        Returns:
            Objeto Document do python-docx.
            
        Raises:
            ArquivoNaoEncontradoError: Se o template não for encontrado.
            TemplateError: Se ocorrer erro ao carregar o template.
        """
        # Atualiza caminho no repositório se fornecido
        if caminho and hasattr(self.repo, 'template_path'):
            self.repo.template_path = caminho
        try:
            caminho_corrente = getattr(self.repo, 'template_path', None)
            logger.info(f"Carregando template: {caminho_corrente}")
            self.documento = self.repo.load()
            logger.info("Template carregado com sucesso")
            return self.documento
        except Exception as e:
            logger.error(f"Erro ao carregar template: {e}")
            raise TemplateError(f"Erro ao carregar template: {e}")
    
    def identificar_placeholders(self):
        """
        Identifica todos os placeholders no documento.
        
        Returns:
            Conjunto com os nomes dos placeholders encontrados.
        """
        if not self.documento:
            self.carregar_template()
        
        self.placeholders_encontrados = set()
        pattern = r'\{\{([^}]+)\}\}'
        
        # Busca em parágrafos
        for paragrafo in self.documento.paragraphs:
            for match in re.finditer(pattern, paragrafo.text):
                ph = match.group(1).strip()
                # Normaliza removendo espaços internos
                ph_norm = ph.replace(' ', '')
                self.placeholders_encontrados.add(ph_norm)
        
        # Busca em tabelas
        for tabela in self.documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for match in re.finditer(pattern, paragrafo.text):
                            ph = match.group(1).strip()
                            # Normaliza removendo espaços internos
                            ph_norm = ph.replace(' ', '')
                            self.placeholders_encontrados.add(ph_norm)
        
        logger.info(f"Placeholders encontrados: {len(self.placeholders_encontrados)}")
        return self.placeholders_encontrados
    
    def validar_placeholders(self):
        """
        Valida se os placeholders encontrados no template têm metadata e vice-versa.
        """
        # Garante que placeholders_encontrados está populado
        if not self.placeholders_encontrados:
            self.identificar_placeholders()
        placeholders_metadata = set(self.metadata.get_all_placeholders())
        missing_in_metadata = self.placeholders_encontrados - placeholders_metadata
        missing_in_template = placeholders_metadata - self.placeholders_encontrados
        if missing_in_metadata:
            logger.warning(f"Placeholders sem metadata encontrados: {missing_in_metadata}")
        if missing_in_template:
            logger.warning(f"Placeholders com metadata mas não usados no template: {missing_in_template}")
        # Em modo estrito, tratar como erro
        if config.MODO_ESTRITO and (missing_in_metadata or missing_in_template):
            raise TemplateError(
                f"Incompatibilidade entre placeholders do template e metadata: "
                f"sem metadata: {missing_in_metadata}, sem uso: {missing_in_template}"
            )
    
    def substituir_placeholders(self, dados, secoes_ativas=None):
        """
        Substitui os placeholders no documento pelos valores dos dados.
        
        Args:
            dados: Dicionário com os valores para substituição.
            secoes_ativas: Lista de IDs das seções que devem estar ativas.
            
        Returns:
            Documento Word com os placeholders substituídos.
            
        Raises:
            SubstituicaoError: Se ocorrer erro na substituição.
        """
        if not self.documento:
            self.carregar_template()
        
        if not secoes_ativas:
            secoes_ativas = []
        
        try:
            # Identifica placeholders se ainda não identificou
            if not self.placeholders_encontrados:
                self.identificar_placeholders()
            
            # Valida placeholders antes da substituição
            self.validar_placeholders()
            
            # Processa parágrafos
            self._substituir_em_paragrafos(dados, secoes_ativas)
            
            # Processa tabelas
            self._substituir_em_tabelas(dados, secoes_ativas)
            
            logger.info("Substituição de placeholders concluída com sucesso")
            return self.documento
        except Exception as e:
            logger.error(f"Erro ao substituir placeholders: {str(e)}")
            raise SubstituicaoError(f"Erro ao substituir placeholders: {str(e)}")
    
    def _substituir_em_paragrafos(self, dados, secoes_ativas):
        """
        Substitui placeholders em parágrafos.
        
        Args:
            dados: Dicionário com os valores para substituição.
            secoes_ativas: Lista de IDs das seções que devem estar ativas.
        """
        pattern = r'\{\{([^}]+)\}\}'
        paragrafos_substituidos = 0
        
        for i, paragrafo in enumerate(self.documento.paragraphs):
            # Verifica se o parágrafo contém algum placeholder
            if not re.search(pattern, paragrafo.text):
                continue
            
            # Verifica se o parágrafo pertence a uma seção inativa
            secao_id = self._identificar_secao_paragrafo(paragrafo.text)
            if secao_id and secao_id not in secoes_ativas:
                # Remove parágrafo de seção inativa
                paragrafo.text = ""
                continue
            
            # Substitui os placeholders
            texto_original = paragrafo.text
            for match in re.finditer(pattern, texto_original):
                placeholder = match.group(1).strip()
                valor_subst = self._obter_valor_substituicao(placeholder, dados)
                
                # Substitui no texto
                paragrafo.text = paragrafo.text.replace(f"{{{{{placeholder}}}}}", str(valor_subst))
            
            if paragrafo.text != texto_original:
                paragrafos_substituidos += 1
        
        logger.debug(f"Parágrafos processados: {paragrafos_substituidos}")
    
    def _substituir_em_tabelas(self, dados, secoes_ativas):
        """
        Substitui placeholders em tabelas.
        
        Args:
            dados: Dicionário com os valores para substituição.
            secoes_ativas: Lista de IDs das seções que devem estar ativas.
        """
        pattern = r'\{\{([^}]+)\}\}'
        celulas_substituidas = 0
        
        for tabela in self.documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        # Verifica se o parágrafo contém algum placeholder
                        if not re.search(pattern, paragrafo.text):
                            continue
                        
                        # Verifica se o parágrafo pertence a uma seção inativa
                        secao_id = self._identificar_secao_paragrafo(paragrafo.text)
                        if secao_id and secao_id not in secoes_ativas:
                            # Remove parágrafo de seção inativa
                            paragrafo.text = ""
                            continue
                        
                        # Substitui os placeholders
                        texto_original = paragrafo.text
                        for match in re.finditer(pattern, texto_original):
                            placeholder = match.group(1).strip()
                            valor_subst = self._obter_valor_substituicao(placeholder, dados)
                            
                            # Substitui no texto
                            paragrafo.text = paragrafo.text.replace(f"{{{{{placeholder}}}}}", str(valor_subst))
                        
                        if paragrafo.text != texto_original:
                            celulas_substituidas += 1
        
        logger.debug(f"Células de tabelas processadas: {celulas_substituidas}")
    
    def _identificar_secao_paragrafo(self, texto):
        """
        Identifica se um parágrafo pertence a uma seção específica.
        No MVP, usamos uma abordagem simples baseada em comentários no texto.
        
        Args:
            texto: Texto do parágrafo.
            
        Returns:
            ID da seção ou None se não for encontrado.
        """
        # Padrão para identificar seções: <!-- SECAO: ID_SECAO -->
        match = re.search(r'<!--\s*SECAO:\s*([A-Za-z0-9_-]+)\s*-->', texto)
        if match:
            return match.group(1)
        return None
    
    def _obter_valor_substituicao(self, placeholder, dados):
        """
        Obtém o valor para substituição de um placeholder.
        
        Args:
            placeholder: Nome do placeholder.
            dados: Dicionário com os valores.
            
        Returns:
            Valor para substituição.
        """
        # Mapeia placeholder -> nome de campo via metadata
        field_name = self.metadata.get_field_name(placeholder)
        if field_name:
            if field_name in dados:
                return self._formatar_valor(dados[field_name])
            else:
                logger.warning(f"Campo '{field_name}' para placeholder '{placeholder}' não presente em dados")
                return ""
        # Placeholder não definido no d_template.csv
        logger.warning(f"Placeholder não definido no d_template: {placeholder}")
        return f"{{DADO NÃO ENCONTRADO: {placeholder}}}"
    
    def _formatar_valor(self, valor):
        """
        Formata um valor de acordo com seu tipo para ser incluído no documento.
        
        Args:
            valor: Valor a ser formatado.
            
        Returns:
            Valor formatado como string.
        """
        # Trata valores None
        if valor is None:
            return ""
            
        # Formata datas (se for objeto datetime)
        if isinstance(valor, datetime):
            return valor.strftime("%d/%m/%Y")
        
        # Verifica se é uma string que parece uma data no formato dd.mm.aaaa
        if isinstance(valor, str):
            # Conversão de datas no formato dd.mm.aaaa
            match = re.match(r'([0-9]{2})\.([0-9]{2})\.([0-9]{4})', valor)
            if match:
                return f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
                
            # Tenta converter números formatados com vírgula decimal
            if re.match(r'^[0-9]+,[0-9]+$', valor):
                try:
                    valor_numerico = float(valor.replace(',', '.'))
                    return f"{valor_numerico:.2f}".replace('.', ',')
                except:
                    pass
        
        # Formata números de forma padronizada
        if isinstance(valor, (int, float)):
            # Formata como decimal com 2 casas se for float
            if isinstance(valor, float):
                return f"{valor:.2f}".replace('.', ',')
            return str(valor)
        
        # Para outros tipos, converte para string
        return str(valor)
    
    def salvar_documento(self, caminho_saida=None):
        """
        Salva o documento processado.
        
        Args:
            caminho_saida: Caminho para o arquivo de saída. Se None, gera um nome no diretório config.OUTPUT_DIR.
            
        Returns:
            Caminho do arquivo salvo.
        """
        if not self.documento:
            logger.error("Nenhum documento carregado para salvar")
            raise TemplateError("Nenhum documento carregado para salvar")
        try:
            caminho = self.repo.save(self.documento, output_name=caminho_saida)
            logger.info(f"Documento salvo: {caminho}")
            return caminho
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}")
            raise SubstituicaoError(f"Erro ao salvar documento: {e}") 