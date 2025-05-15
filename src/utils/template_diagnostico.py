"""
Ferramenta de diagnóstico para templates DOCX.

Esta ferramenta examina documentos DOCX e identifica detalhadamente os placeholders,
incluindo aqueles potencialmente fragmentados em várias "runs" ou com problemas de formatação.
"""

import os
import sys
import re
import docx
import argparse
from typing import List, Dict, Set, Any, Tuple
from collections import defaultdict

# Adiciona o diretório pai ao path para importar módulos
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from src.logger import logger

class TemplateDiagnostico:
    """
    Ferramenta para diagnóstico de templates do sistema de peticionamento.
    """
    
    def __init__(self, modo_verbose: bool = False):
        """
        Inicializa o diagnóstico de templates.
        
        Args:
            modo_verbose: Se True, exibe informações detalhadas durante o processamento.
        """
        self.modo_verbose = modo_verbose
        self.placeholders_detectados = []
        self.stats = {
            "total_paragrafos": 0,
            "total_runs": 0,
            "total_placeholders": 0,
            "placeholders_simples": 0,
            "placeholders_fragmentados": 0,
            "placeholders_secao_inicio": 0,
            "placeholders_secao_fim": 0,
            "placeholders_malformados": 0,
            "textos_estranhos": 0
        }
    
    def analisar_template(self, template_path: str) -> Dict[str, Any]:
        """
        Analisa um template DOCX e identifica placeholders e problemas.
        
        Args:
            template_path: Caminho para o arquivo de template DOCX.
            
        Returns:
            Dicionário com estatísticas e problemas encontrados.
        """
        if not os.path.exists(template_path):
            print(f"Erro: Template não encontrado: {template_path}")
            return {"erro": "Template não encontrado"}
        
        print(f"\n{'='*50}")
        print(f"DIAGNÓSTICO DO TEMPLATE: {os.path.basename(template_path)}")
        print(f"{'='*50}\n")
        
        # Abre o documento
        try:
            doc = docx.Document(template_path)
        except Exception as e:
            print(f"Erro ao abrir o documento: {str(e)}")
            return {"erro": f"Erro ao abrir o documento: {str(e)}"}
        
        self.stats["total_paragrafos"] = len(doc.paragraphs)
        
        # 1. Análise básica - paragráfos e runs
        print("1. ANÁLISE BÁSICA DE PARÁGRAFOS E RUNS")
        print(f"   Total de parágrafos: {len(doc.paragraphs)}")
        
        total_runs = sum(len(p.runs) for p in doc.paragraphs)
        self.stats["total_runs"] = total_runs
        print(f"   Total de runs: {total_runs}")
        
        # 2. Busca por placeholders em texto normal
        print("\n2. BUSCA POR PLACEHOLDERS EM TEXTO NORMAL")
        self._analisar_placeholders_simples(doc)
        
        # 3. Busca por placeholders fragmentados
        print("\n3. BUSCA POR PLACEHOLDERS FRAGMENTADOS")
        self._analisar_placeholders_fragmentados(doc)
        
        # 4. Busca por seções condicionais
        print("\n4. BUSCA POR SEÇÕES CONDICIONAIS")
        self._analisar_secoes_condicionais(doc)
        
        # 5. Busca por problemas e anomalias
        print("\n5. BUSCA POR PROBLEMAS E ANOMALIAS")
        self._analisar_problemas(doc)
        
        # 6. Resumo estatístico
        print("\n6. RESUMO ESTATÍSTICO")
        self._exibir_estatisticas()
        
        # 7. Recomendações
        print("\n7. RECOMENDAÇÕES PARA CORREÇÃO")
        self._gerar_recomendacoes()
        
        return {
            "estatisticas": self.stats,
            "placeholders": self.placeholders_detectados
        }
    
    def _analisar_placeholders_simples(self, doc: docx.Document) -> None:
        """
        Busca placeholders em parágrafos de texto normal (não fragmentados).
        
        Args:
            doc: Documento DOCX a ser analisado.
        """
        pattern = r'{{[\s]*([^{}]+?)[\s]*}}'
        placeholders_simples = 0
        
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text
            matches = re.finditer(pattern, texto)
            
            for match in matches:
                placeholder = match.group(0)
                campo = match.group(1).strip()
                
                # Ignora marcadores de seção, serão analisados separadamente
                if campo.startswith('#') or campo.startswith('/'):
                    continue
                
                placeholders_simples += 1
                self.placeholders_detectados.append({
                    "tipo": "simples",
                    "placeholder": placeholder,
                    "campo": campo,
                    "paragrafo": i+1,
                    "texto_paragrafo": texto[:50] + "..." if len(texto) > 50 else texto
                })
                
                if self.modo_verbose:
                    print(f"   Placeholder: '{placeholder}' (Campo: '{campo}')")
                    print(f"   Localização: Parágrafo {i+1}")
                    print(f"   Texto: '{texto[:50]}{'...' if len(texto) > 50 else ''}'")
                    print("")
        
        self.stats["placeholders_simples"] = placeholders_simples
        self.stats["total_placeholders"] += placeholders_simples
        
        print(f"   Total de placeholders simples encontrados: {placeholders_simples}")
    
    def _analisar_placeholders_fragmentados(self, doc: docx.Document) -> None:
        """
        Busca placeholders que podem estar fragmentados entre diferentes runs.
        
        Args:
            doc: Documento DOCX a ser analisado.
        """
        placeholders_fragmentados = 0
        
        for i, paragrafo in enumerate(doc.paragraphs):
            if len(paragrafo.runs) <= 1:
                continue
            
            # Reconstruir o texto concatenado das runs
            runs_texto = [run.text for run in paragrafo.runs]
            texto_completo = "".join(runs_texto)
            
            # Buscar placeholders no texto concatenado
            pattern = r'{{[\s]*([^{}]+?)[\s]*}}'
            matches = list(re.finditer(pattern, texto_completo))
            
            # Se não encontrar nada no texto concatenado, continua para o próximo parágrafo
            if not matches:
                continue
            
            # Verificar se algum placeholder está fragmentado
            for match in matches:
                placeholder = match.group(0)
                campo = match.group(1).strip()
                inicio = match.start()
                fim = match.end()
                
                # Ignora marcadores de seção, serão analisados separadamente
                if campo.startswith('#') or campo.startswith('/'):
                    continue
                
                # Determina em quais runs o placeholder está
                runs_afetadas = []
                indices_em_runs = []
                posicao_atual = 0
                
                for j, texto_run in enumerate(runs_texto):
                    nova_posicao = posicao_atual + len(texto_run)
                    
                    # Run contém parte do placeholder
                    if (posicao_atual <= inicio < nova_posicao) or \
                       (posicao_atual < fim <= nova_posicao) or \
                       (posicao_atual >= inicio and nova_posicao <= fim):
                        runs_afetadas.append(j)
                    
                    posicao_atual = nova_posicao
                
                # Se o placeholder está em mais de um run, é fragmentado
                if len(runs_afetadas) > 1:
                    placeholders_fragmentados += 1
                    
                    # Registra o placeholder fragmentado
                    self.placeholders_detectados.append({
                        "tipo": "fragmentado",
                        "placeholder": placeholder,
                        "campo": campo,
                        "paragrafo": i+1,
                        "runs_afetadas": runs_afetadas,
                        "texto_paragrafo": texto_completo[:50] + "..." if len(texto_completo) > 50 else texto_completo
                    })
                    
                    # Exibe detalhes
                    print(f"   Placeholder fragmentado: '{placeholder}' (Campo: '{campo}')")
                    print(f"   Localização: Parágrafo {i+1}, Runs {runs_afetadas}")
                    
                    if self.modo_verbose:
                        print("   Detalhes das runs:")
                        for j in runs_afetadas:
                            print(f"     Run {j}: '{paragrafo.runs[j].text}'")
                    
                    print("")
        
        self.stats["placeholders_fragmentados"] = placeholders_fragmentados
        self.stats["total_placeholders"] += placeholders_fragmentados
        
        print(f"   Total de placeholders fragmentados encontrados: {placeholders_fragmentados}")
    
    def _analisar_secoes_condicionais(self, doc: docx.Document) -> None:
        """
        Busca marcadores de seções condicionais no documento.
        
        Args:
            doc: Documento DOCX a ser analisado.
        """
        secoes_inicio = 0
        secoes_fim = 0
        secoes_dict = {}
        
        inicio_pattern = r'{{[\s]*#[\s]*SECAO[\s_]*([A-Za-z0-9_]+)[\s]*}}'
        fim_pattern = r'{{[\s]*/[\s]*SECAO[\s_]*([A-Za-z0-9_]+)[\s]*}}'
        
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text
            
            # Busca início de seção
            match_inicio = re.search(inicio_pattern, texto, re.IGNORECASE)
            if match_inicio:
                secao_id = match_inicio.group(1).upper()
                secoes_inicio += 1
                
                if secao_id not in secoes_dict:
                    secoes_dict[secao_id] = {"inicio": [], "fim": []}
                
                secoes_dict[secao_id]["inicio"].append(i+1)
                
                self.placeholders_detectados.append({
                    "tipo": "secao_inicio",
                    "secao_id": secao_id,
                    "placeholder": match_inicio.group(0),
                    "paragrafo": i+1,
                    "texto_paragrafo": texto[:50] + "..." if len(texto) > 50 else texto
                })
                
                print(f"   Início de seção: '{match_inicio.group(0)}' (ID: '{secao_id}')")
                print(f"   Localização: Parágrafo {i+1}")
                print("")
            
            # Busca fim de seção
            match_fim = re.search(fim_pattern, texto, re.IGNORECASE)
            if match_fim:
                secao_id = match_fim.group(1).upper()
                secoes_fim += 1
                
                if secao_id not in secoes_dict:
                    secoes_dict[secao_id] = {"inicio": [], "fim": []}
                
                secoes_dict[secao_id]["fim"].append(i+1)
                
                self.placeholders_detectados.append({
                    "tipo": "secao_fim",
                    "secao_id": secao_id,
                    "placeholder": match_fim.group(0),
                    "paragrafo": i+1,
                    "texto_paragrafo": texto[:50] + "..." if len(texto) > 50 else texto
                })
                
                print(f"   Fim de seção: '{match_fim.group(0)}' (ID: '{secao_id}')")
                print(f"   Localização: Parágrafo {i+1}")
                print("")
        
        # Verifica consistência de seções
        print("\n   Consistência de seções:")
        for secao_id, info in secoes_dict.items():
            if len(info["inicio"]) > 0 and len(info["fim"]) > 0:
                print(f"   ✓ Seção '{secao_id}': {len(info['inicio'])} início(s) e {len(info['fim'])} fim(s)")
            else:
                if len(info["inicio"]) == 0:
                    print(f"   ✗ Seção '{secao_id}': Falta marcador de início!")
                if len(info["fim"]) == 0:
                    print(f"   ✗ Seção '{secao_id}': Falta marcador de fim!")
        
        self.stats["placeholders_secao_inicio"] = secoes_inicio
        self.stats["placeholders_secao_fim"] = secoes_fim
        self.stats["total_placeholders"] += secoes_inicio + secoes_fim
        
        print(f"\n   Total de marcadores de início de seção: {secoes_inicio}")
        print(f"   Total de marcadores de fim de seção: {secoes_fim}")
    
    def _analisar_problemas(self, doc: docx.Document) -> None:
        """
        Busca problemas e anomalias no documento.
        
        Args:
            doc: Documento DOCX a ser analisado.
        """
        problemas_detectados = 0
        
        # 1. Busca por placeholders malformados
        malformados_pattern = r'{{[^}]*$|^[^{]*}}'
        placeholders_malformados = 0
        
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text
            
            # Verifica chaves não fechadas ou não abertas
            if re.search(malformados_pattern, texto):
                placeholders_malformados += 1
                print(f"   ✗ Placeholder malformado no parágrafo {i+1}: '{texto[:50]}{'...' if len(texto) > 50 else ''}'")
                
                self.placeholders_detectados.append({
                    "tipo": "malformado",
                    "paragrafo": i+1,
                    "texto_paragrafo": texto[:50] + "..." if len(texto) > 50 else texto
                })
        
        self.stats["placeholders_malformados"] = placeholders_malformados
        problemas_detectados += placeholders_malformados
        
        # 2. Busca por caracteres estranhos ou invisíveis
        textos_estranhos = 0
        
        for i, paragrafo in enumerate(doc.paragraphs):
            texto = paragrafo.text
            
            # Verifica caracteres de controle invisíveis (exceto espaços e tabs)
            if re.search(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', texto):
                textos_estranhos += 1
                
                # Converte para hexadecimal para visualização
                texto_hex = ' '.join(f'{ord(c):02x}' for c in texto)
                
                print(f"   ✗ Caracteres invisíveis no parágrafo {i+1}")
                print(f"     Texto Hex: {texto_hex[:100]}{'...' if len(texto_hex) > 100 else ''}")
                
                self.placeholders_detectados.append({
                    "tipo": "caracteres_estranhos",
                    "paragrafo": i+1,
                    "texto_paragrafo": texto[:50] + "..." if len(texto) > 50 else texto,
                    "texto_hex": texto_hex
                })
        
        self.stats["textos_estranhos"] = textos_estranhos
        problemas_detectados += textos_estranhos
        
        print(f"\n   Total de problemas detectados: {problemas_detectados}")
    
    def _exibir_estatisticas(self) -> None:
        """
        Exibe um resumo estatístico da análise.
        """
        print(f"   Total de parágrafos: {self.stats['total_paragrafos']}")
        print(f"   Total de runs: {self.stats['total_runs']}")
        print(f"   Total de placeholders: {self.stats['total_placeholders']}")
        print(f"     - Placeholders simples: {self.stats['placeholders_simples']}")
        print(f"     - Placeholders fragmentados: {self.stats['placeholders_fragmentados']}")
        print(f"     - Marcadores de início de seção: {self.stats['placeholders_secao_inicio']}")
        print(f"     - Marcadores de fim de seção: {self.stats['placeholders_secao_fim']}")
        print(f"   Total de problemas: {self.stats['placeholders_malformados'] + self.stats['textos_estranhos']}")
        print(f"     - Placeholders malformados: {self.stats['placeholders_malformados']}")
        print(f"     - Textos com caracteres estranhos: {self.stats['textos_estranhos']}")
    
    def _gerar_recomendacoes(self) -> None:
        """
        Gera recomendações com base nos problemas encontrados.
        """
        has_recommendations = False
        
        # Problemas com seções
        if self.stats['placeholders_secao_inicio'] == 0 and self.stats['placeholders_secao_fim'] == 0:
            has_recommendations = True
            print("   1. O template não contém nenhuma seção condicional.")
            print("      Recomendação: Adicione seções condicionais usando os marcadores {{#SECAO_NOME}} e {{/SECAO_NOME}}.")
            print("      Exemplo: {{#SECAO_PEDIDO_LIMINAR}} ... texto condicional ... {{/SECAO_PEDIDO_LIMINAR}}")
        elif self.stats['placeholders_secao_inicio'] != self.stats['placeholders_secao_fim']:
            has_recommendations = True
            print("   2. O número de marcadores de início de seção não corresponde ao número de marcadores de fim.")
            print("      Recomendação: Verifique se cada seção tem corretamente um marcador de início e um de fim.")
            print("      Corrija os marcadores ausentes ou em excesso.")
        
        # Problemas com placeholders fragmentados
        if self.stats['placeholders_fragmentados'] > 0:
            has_recommendations = True
            print("   3. Existem placeholders fragmentados entre múltiplas runs.")
            print("      Recomendação: Recrie esses placeholders mantendo-os em uma única run:")
            print("      a) Selecione todo o texto do placeholder.")
            print("      b) Recorte (Ctrl+X) e cole novamente (Ctrl+V) para unificar a formatação.")
            print("      c) Edite o texto para garantir que esteja no formato correto: {{nome_campo}}")
        
        # Problemas com placeholders malformados
        if self.stats['placeholders_malformados'] > 0:
            has_recommendations = True
            print("   4. Existem placeholders malformados (chaves não fechadas ou não abertas).")
            print("      Recomendação: Corrija esses placeholders para seguir o formato padrão {{nome_campo}}.")
            print("      Verifique especialmente os parágrafos indicados acima.")
        
        # Caracteres estranhos
        if self.stats['textos_estranhos'] > 0:
            has_recommendations = True
            print("   5. Existem caracteres invisíveis ou de controle em alguns parágrafos.")
            print("      Recomendação: Recrie esses parágrafos com texto limpo para evitar problemas.")
            print("      a) Crie um novo parágrafo vazio.")
            print("      b) Digite novamente o texto, incluindo os placeholders.")
        
        # Uso do conversor de templates
        has_recommendations = True
        print("   6. Para corrigir automaticamente alguns dos problemas acima, use a ferramenta de conversão de templates:")
        print("      python src/utils/template_converter.py template.docx template_corrigido.docx")
        
        if not has_recommendations:
            print("   ✓ Não foram detectados problemas significativos no template.")
            print("     O template parece estar em boas condições para processamento.")


def main():
    """
    Função principal para execução direta do script.
    """
    parser = argparse.ArgumentParser(description="Diagnóstico detalhado de templates DOCX")
    parser.add_argument("template", help="Caminho para o arquivo de template DOCX")
    parser.add_argument("--verbose", "-v", action="store_true", help="Exibe informações detalhadas durante o processamento")
    
    args = parser.parse_args()
    
    diagnostico = TemplateDiagnostico(modo_verbose=args.verbose)
    diagnostico.analisar_template(args.template)


if __name__ == "__main__":
    main() 