"""
Utilitário para converter templates para o formato padronizado de seções e corrigir placeholders fragmentados.
"""

import os
import re
import sys
import docx
import logging
from typing import Dict, List, Tuple, Any, Set, Optional, Union, cast
from docx.document import Document
from docx.enum.text import WD_UNDERLINE

# Adiciona o diretório pai ao path para importar módulos
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from src.utils.logger_config import configurar_logger

logger = configurar_logger()

def converter_template_para_formato_padrao(template_path: str, output_path: str, secoes_conhecidas: Optional[Dict[str, Dict[str, str]]] = None) -> bool:
    """
    Converte um template existente para usar o formato padronizado de seções e corrige placeholders fragmentados.
    
    Args:
        template_path: Caminho do template a ser convertido
        output_path: Caminho onde o template convertido será salvo
        secoes_conhecidas: Dicionário com definições de seções conhecidas
                         Formato: {
                             "SECAO_ID": {
                                 "inicio": "Texto que marca o início da seção",
                                 "fim": "Texto que marca o fim da seção"
                             }
                         }
                         
    Returns:
        True se a conversão for bem-sucedida, False caso contrário
    """
    if secoes_conhecidas is None:
        secoes_conhecidas = {
            "HORAS_EXTRAS": {
                "inicio": "HORAS EXTRAS",
                "fim": "VALOR REQUERIDO"
            },
            "VERBAS_RESCISORIAS": {
                "inicio": "VERBAS RESCISÓRIAS",
                "fim": "CONCLUSÃO"
            },
            "INSALUBRIDADE": {
                "inicio": "INSALUBRIDADE",
                "fim": "VALOR REQUERIDO"
            },
            "DANO_MORAL": {
                "inicio": "DANO MORAL",
                "fim": "VALOR REQUERIDO"
            },
            "ACUMULO_FUNCAO": {
                "inicio": "ACÚMULO DE FUNÇÃO",
                "fim": "VALOR REQUERIDO"
            },
            "AVISO_PREVIO": {
                "inicio": "AVISO PRÉVIO",
                "fim": "VERBAS RESCISÓRIAS"
            }
        }
    
    try:
        # Verifica se o template existe
        if not os.path.exists(template_path):
            logger.error(f"Template não encontrado: {template_path}")
            return False
        
        # Carrega o documento
        doc = docx.Document(template_path)
        logger.info(f"Template carregado: {template_path}")
        
        # Corrige placeholders fragmentados
        placeholders_corrigidos = corrigir_placeholders_fragmentados(doc)
        logger.info(f"Placeholders corrigidos: {placeholders_corrigidos}")
        
        # Busca por placeholders malformados ou texto com problemas
        placeholders_malformados_corrigidos = corrigir_placeholders_malformados(doc)
        logger.info(f"Placeholders malformados corrigidos: {placeholders_malformados_corrigidos}")
        
        # Mapeia posições de início e fim de cada seção
        secoes_mapeadas = mapear_secoes(doc, secoes_conhecidas)
        logger.info(f"Seções mapeadas: {len(secoes_mapeadas)}")
        
        # Insere marcadores de seção
        if secoes_mapeadas:
            inserir_marcadores_secao(doc, secoes_mapeadas)
            logger.info(f"Marcadores de seção inseridos")
        else:
            logger.warning("Nenhuma seção mapeada, apenas placeholders foram corrigidos")
        
        # Salva o documento modificado
        doc.save(output_path)
        logger.info(f"Template convertido salvo em: {output_path}")
        return True
            
    except Exception as e:
        logger.error(f"Erro ao converter template: {str(e)}")
        return False

def mapear_secoes(doc: Document, secoes_conhecidas: Dict[str, Dict[str, str]]) -> Dict[str, Dict[str, int]]:
    """
    Mapeia posições de início e fim de cada seção no documento.
    
    Args:
        doc: Documento docx
        secoes_conhecidas: Definições de seções conhecidas
        
    Returns:
        Dicionário com informações sobre cada seção mapeada
    """
    secoes_mapeadas: Dict[str, Dict[str, int]] = {}
    
    # Primeiro, verifica se já existem marcadores de seção no documento
    secoes_existentes = identificar_secoes_existentes(doc)
    if secoes_existentes:
        logger.info(f"Encontradas {len(secoes_existentes)} seções já marcadas no documento")
        return {}  # Retorna vazio para não inserir marcadores duplicados
    
    # Procura textos de início e fim de cada seção
    for i, paragrafo in enumerate(doc.paragraphs):
        texto = paragrafo.text.strip()
        
        # Verifica se o texto corresponde ao início de alguma seção conhecida
        for secao_id, info in secoes_conhecidas.items():
            texto_inicio = info.get("inicio", "")
            
            if texto_inicio and texto_inicio in texto:
                logger.debug(f"Encontrado início da seção {secao_id} no parágrafo {i}: '{texto}'")
                if secao_id not in secoes_mapeadas:
                    secoes_mapeadas[secao_id] = {"inicio": i}
            
            texto_fim = info.get("fim", "")
            if texto_fim and texto_fim in texto:
                logger.debug(f"Encontrado fim da seção {secao_id} no parágrafo {i}: '{texto}'")
                if secao_id in secoes_mapeadas and "fim" not in secoes_mapeadas[secao_id]:
                    secoes_mapeadas[secao_id]["fim"] = i
    
    # Verifica quais seções têm início e fim definidos
    secoes_completas: Dict[str, Dict[str, int]] = {}
    for secao_id, posicoes in secoes_mapeadas.items():
        if "inicio" in posicoes and "fim" in posicoes:
            inicio = posicoes["inicio"]
            fim = posicoes["fim"]
            
            # Só mantém se o fim vem depois do início
            if fim > inicio:
                secoes_completas[secao_id] = posicoes
                logger.info(f"Seção {secao_id} mapeada: parágrafos {inicio} a {fim}")
            else:
                logger.warning(f"Seção {secao_id} tem fim ({fim}) antes do início ({inicio}). Ignorando.")
        else:
            logger.warning(f"Seção {secao_id} não tem início e fim definidos. Ignorando.")
    
    return secoes_completas

def identificar_secoes_existentes(doc: Document) -> Dict[str, Dict[str, int]]:
    """
    Identifica seções já marcadas no documento.
    
    Args:
        doc: Documento docx
        
    Returns:
        Dicionário com informações sobre cada seção já marcada
    """
    secoes: Dict[str, Dict[str, int]] = {}
    inicio_pattern = r'{{[\s]*#[\s]*SECAO[\s_]*([A-Za-z0-9_]+)[\s]*}}'
    fim_pattern = r'{{[\s]*/[\s]*SECAO[\s_]*([A-Za-z0-9_]+)[\s]*}}'
    
    for i, paragrafo in enumerate(doc.paragraphs):
        texto = paragrafo.text.strip()
        
        # Busca início de seção
        match_inicio = re.search(inicio_pattern, texto, re.IGNORECASE)
        if match_inicio:
            secao_id = match_inicio.group(1).upper()
            if secao_id not in secoes:
                secoes[secao_id] = {"inicio": i}
        
        # Busca fim de seção
        match_fim = re.search(fim_pattern, texto, re.IGNORECASE)
        if match_fim:
            secao_id = match_fim.group(1).upper()
            if secao_id in secoes and "fim" not in secoes[secao_id]:
                secoes[secao_id]["fim"] = i
    
    return secoes

def inserir_marcadores_secao(doc: Document, secoes_mapeadas: Dict[str, Dict[str, int]]) -> None:
    """
    Insere marcadores de início e fim de seção no documento.
    
    Args:
        doc: Documento docx
        secoes_mapeadas: Informações sobre as seções mapeadas
    """
    # Processa seções de trás para frente para evitar que as inserções
    # alterem os índices das seções subsequentes
    for secao_id, posicoes in sorted(secoes_mapeadas.items(), key=lambda x: x[1]["inicio"], reverse=True):
        inicio = posicoes["inicio"]
        fim = posicoes["fim"]
        
        # Insere marcador de fim após o parágrafo de fim
        paragrafo_fim = doc.paragraphs[fim]
        run = paragrafo_fim.add_run()
        run.add_break()
        paragrafo_fim_marcador = doc.add_paragraph(f"{{{{/SECAO_{secao_id}}}}}")
        
        # Insere marcador de início antes do parágrafo de início
        paragrafo_inicio = doc.paragraphs[inicio]
        paragrafo_inicio_marcador = paragrafo_inicio.insert_paragraph_before(f"{{{{#SECAO_{secao_id}}}}}")

def corrigir_placeholders_fragmentados(doc: Document) -> int:
    """
    Detecta e corrige placeholders fragmentados em diferentes runs.
    
    Args:
        doc: Documento docx
        
    Returns:
        Número de placeholders corrigidos
    """
    total_corrigidos = 0
    
    for i, paragrafo in enumerate(doc.paragraphs):
        # Se o parágrafo tem apenas uma run, não precisa verificar fragmentação
        if len(paragrafo.runs) <= 1:
            continue
        
        # Reconstrói o texto completo do parágrafo concatenando todas as runs
        runs_texto = [run.text for run in paragrafo.runs]
        texto_completo = "".join(runs_texto)
        
        # Busca por placeholders no texto completo
        placeholders = list(re.finditer(r'{{[\s]*([^{}]+?)[\s]*}}', texto_completo))
        
        # Se não encontrar placeholders, pula para o próximo parágrafo
        if not placeholders:
            continue
        
        # Verifica se algum placeholder está fragmentado
        tem_fragmentado = False
        for match in placeholders:
            inicio = match.start()
            fim = match.end()
            placeholder = match.group(0)
            
            # Determina em quais runs o placeholder está
            runs_afetadas = []
            posicao_atual = 0
            
            for j, texto_run in enumerate(runs_texto):
                nova_posicao = posicao_atual + len(texto_run)
                
                # Run contém parte do placeholder
                if (posicao_atual <= inicio < nova_posicao) or \
                   (posicao_atual < fim <= nova_posicao) or \
                   (posicao_atual >= inicio and nova_posicao <= fim):
                    runs_afetadas.append(j)
                
                posicao_atual = nova_posicao
            
            # Se o placeholder está em mais de um run, está fragmentado
            if len(runs_afetadas) > 1:
                tem_fragmentado = True
                logger.debug(f"Placeholder fragmentado no parágrafo {i+1}: '{placeholder}'")
        
        # Se o parágrafo tem placeholders fragmentados, substitui o texto para corrigir
        if tem_fragmentado:
            # Guarda a formatação da primeira run para aplicar no novo texto
            run_modelo = paragrafo.runs[0]
            
            # Limpa o parágrafo atual
            for _ in range(len(paragrafo.runs)):
                if paragrafo.runs and paragrafo._p and paragrafo.runs[0]._r:
                    paragrafo._p.remove(paragrafo.runs[0]._r)
            
            # Adiciona uma nova run com o texto completo e a formatação original
            nova_run = paragrafo.add_run(texto_completo)
            
            # Aplica a formatação da run modelo
            nova_run.bold = run_modelo.bold
            nova_run.italic = run_modelo.italic
            # Tratando especificamente para o caso do underline
            if run_modelo.underline:
                if isinstance(run_modelo.underline, bool):
                    nova_run.underline = run_modelo.underline
                else:
                    # Se não for bool, é uma constante WD_UNDERLINE
                    nova_run.font.underline = run_modelo.underline
                
            if hasattr(run_modelo.font, 'name') and run_modelo.font.name:
                nova_run.font.name = run_modelo.font.name
            if hasattr(run_modelo.font, 'size') and run_modelo.font.size:
                nova_run.font.size = run_modelo.font.size
            
            # Incrementa o contador
            total_corrigidos += len(placeholders)
            logger.info(f"Parágrafo {i+1}: {len(placeholders)} placeholders corrigidos")
    
    return total_corrigidos

def corrigir_placeholders_malformados(doc: Document) -> int:
    """
    Detecta e corrige placeholders malformados (com chaves abertas/fechadas incorretamente).
    
    Args:
        doc: Documento docx
        
    Returns:
        Número de placeholders corrigidos
    """
    total_corrigidos = 0
    
    # Padrão para detectar chaves não fechadas ou não abertas corretamente
    malformados_pattern = r'{{[^}]+$|^[^{]+}}'
    
    for i, paragrafo in enumerate(doc.paragraphs):
        texto = paragrafo.text
        
        # Verifica se há chaves desbalanceadas contando número de { e }
        contagem_abre = texto.count('{')
        contagem_fecha = texto.count('}')
        
        if contagem_abre != contagem_fecha:
            logger.info(f"Parágrafo {i+1} tem chaves desbalanceadas: {contagem_abre} abertas, {contagem_fecha} fechadas")
            
            # Tenta corrigir chaves desbalanceadas
            texto_corrigido = texto
            
            # Se tiver {{ sem }}, adiciona }}
            if texto_corrigido.count('{{') > texto_corrigido.count('}}'):
                texto_corrigido += '}}'
                
            # Se tiver }} sem {{, adiciona {{ no início
            if texto_corrigido.count('}}') > texto_corrigido.count('{{'):
                texto_corrigido = '{{' + texto_corrigido
            
            # Atualiza o texto do parágrafo
            if texto_corrigido != texto:
                # Guarda a formatação
                run_modelo = paragrafo.runs[0] if paragrafo.runs else None
                
                # Limpa o parágrafo atual
                for _ in range(len(paragrafo.runs)):
                    if paragrafo.runs and paragrafo._p and paragrafo.runs[0]._r:
                        paragrafo._p.remove(paragrafo.runs[0]._r)
                
                # Adiciona o texto corrigido
                nova_run = paragrafo.add_run(texto_corrigido)
                
                # Aplica a formatação original, se disponível
                if run_modelo:
                    nova_run.bold = run_modelo.bold
                    nova_run.italic = run_modelo.italic
                    # Tratando especificamente para o caso do underline
                    if run_modelo.underline:
                        if isinstance(run_modelo.underline, bool):
                            nova_run.underline = run_modelo.underline
                        else:
                            # Se não for bool, é uma constante WD_UNDERLINE
                            nova_run.font.underline = run_modelo.underline
                        
                    if hasattr(run_modelo.font, 'name') and run_modelo.font.name:
                        nova_run.font.name = run_modelo.font.name
                    if hasattr(run_modelo.font, 'size') and run_modelo.font.size:
                        nova_run.font.size = run_modelo.font.size
                
                total_corrigidos += 1
                logger.info(f"Parágrafo {i+1}: placeholder malformado corrigido")
    
    return total_corrigidos

if __name__ == "__main__":
    # Exemplo de uso do script diretamente
    import argparse
    
    parser = argparse.ArgumentParser(description="Converte templates para o formato padronizado de seções e corrige placeholders")
    parser.add_argument("template", help="Caminho do template a ser convertido")
    parser.add_argument("output", help="Caminho onde o template convertido será salvo")
    
    args = parser.parse_args()
    
    resultado = converter_template_para_formato_padrao(args.template, args.output)
    if resultado:
        print(f"Template convertido com sucesso! Salvo em: {args.output}")
    else:
        print("Erro ao converter o template. Verifique os logs para mais detalhes.")
        sys.exit(1) 